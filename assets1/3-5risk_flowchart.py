"""Risk Management Governance Structure flowchart for Word reports (3-5).
Creates a flowchart using matplotlib and inserts it into Word."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
from matplotlib.patches import FancyArrowPatch, Rectangle
from io import BytesIO

# Flowchart data
BOXES = [
    (-0.35, 0.78, 0.4, 0.12, '#009999', "Supervisory Board\n• Request to investigate risk topics\n• Bi-annual risk review"),
    (0.35, 0.78, 0.4, 0.12, '#009999', "Audit Committee\n• Assertion on control effectiveness\n• Quarterly progress reporting"),
    (0, 0.55, 0.55, 0.10, '#0073cf', "Board of Management"),
    (-0.35, 0.35, 0.4, 0.12, '#00a9e0', "Security and Risk Committee (CESR)\nRisk oversight\n• Risk appetite\n• Risk management policy"),
    (0.35, 0.35, 0.4, 0.12, '#00a9e0', "Disclosure Committee\nInternal Control Committee\n• Control effectiveness\n• Risk response progress"),
    (0, 0.15, 0.55, 0.10, '#1b3b7a', "Risk Owners")
]

ARROWS = [
    (-0.35, 0.70, 0, 0.60),
    (0.35, 0.70, 0, 0.60),
    (-0.15, 0.50, -0.35, 0.42),
    (0.15, 0.50, 0.35, 0.42),
    (-0.35, 0.29, 0, 0.20),
    (0.35, 0.29, 0, 0.20),
]


def hex_to_rgb(hex_color):
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) / 255.0 for i in (0, 2, 4))


def add_arrow(ax, start, end, color="#1f4e79", lw=4):
    arrow = FancyArrowPatch(posA=start, posB=end, arrowstyle="-|>", color=color, linewidth=lw, mutation_scale=24, connectionstyle="arc3,rad=0.08")
    ax.add_patch(arrow)


def create_flowchart_image():
    # Use a generous canvas, but we'll trim all margins strictly when saving
    fig, ax = plt.subplots(figsize=(11.0, 6.8))
    # Title inside plot area to avoid top white margin
    ax.text(0, 1.02, "Risk Management Governance Structure", ha='left', va='bottom', fontsize=20, fontweight='bold', color='#1f4e79', transform=ax.transAxes)
    for (x, y, w, h, color, text) in BOXES:
        rect = Rectangle((x - w/2, y - h/2), w, h, facecolor=hex_to_rgb(color), edgecolor='white', linewidth=2.0, zorder=1)
        ax.add_patch(rect)
        ax.text(x, y, text, ha='center', va='center', fontsize=11, color='white', fontweight='bold', wrap=True, zorder=2)
    for (sx, sy, ex, ey) in ARROWS:
        add_arrow(ax, (sx, sy), (ex, ey))
    ax.axis('off'); ax.set_xlim(-1, 1); ax.set_ylim(0, 1); ax.set_aspect('equal')
    # Remove all outer margins
    plt.margins(0)
    fig.tight_layout(pad=0)
    ax.set_position([0, 0, 1, 1])
    buf = BytesIO(); plt.savefig(buf, format='png', dpi=300, bbox_inches='tight', pad_inches=0, transparent=True); buf.seek(0); plt.close()
    return buf


def create_table():
    doc = Document()
    title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Risk Management Governance Structure"); run.font.size = Pt(16); run.font.bold = True; run.font.color.rgb = RGBColor(31, 78, 121)
    img = create_flowchart_image()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(img, width=Inches(10))
    return doc


if __name__ == "__main__":
    d = create_table(); d.save("preview_risk_flowchart_3_5.docx")
