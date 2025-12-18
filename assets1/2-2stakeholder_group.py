"""Stakeholder group table for Word reports.
Creates a table with no borders (borderless), only bottom borders for rows."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH

STAKEHOLDER_DATA = [
    {
        "group": "Investors & Shareholders",
        "topics": "Climate Resilience, TCFD Disclosure, Financial Performance, Governance Structure, and ESG-linked Executive Compensation.",
        "channels": "Annual General Meetings (AGM), Quarterly Earnings Calls, Communications with ESG Rating Agencies (MSCI, Sustainalytics), and Corporate Communications.",
    },
    {
        "group": "Employees & Potential Hires",
        "topics": "Talent Attraction/Retention, Diversity, Equity & Inclusion (DEI), Health & Safety, Ethical Culture (Anti-corruption training).",
        "channels": "Annual Employee Pulse Surveys, Internal Training Workshops, Town Halls with Senior Management, and Staff Grievance Procedures.",
    },
    {
        "group": "Customers & Clients",
        "topics": "Data Privacy & Cybersecurity, Fair Treatment (TCF), Product Safety, Service Accessibility, and Sustainable Finance Options (Green Loans).",
        "channels": "Customer Feedback Channels, Customer Service Hotlines, Annual Mystery Shopper Studies, and Net Promoter Score (NPS) Surveys.",
    },
    {
        "group": "Regulators & Governments",
        "topics": "Compliance with laws (HKMA, SFC, GDPR/POPIA), Risk Management Integration (ERM/CRST), and adherence to new standards (ISSB/ESRS).",
        "channels": "Official Regulatory Submissions, Quarterly Compliance Reports, Participation in Industry Forums, and Policy Advocacy.",
    },
    {
        "group": "Suppliers & Partners",
        "topics": "Responsible Supply Chain Management, Human Rights (Anti-Slavery), Ethical Sourcing, and Sustainable Procurement Policy compliance.",
        "channels": "Supplier Code of Conduct (SCoC) Acknowledgement, Third-Party Risk Assessments, and Due Diligence Audits.",
    },
]


def _apply_shading(cell, color_hex: str):
    """Apply background shading to cell"""
    shading = parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading)


def _format_text(cell, text, bold=False, size=10, color=(31, 41, 55)):
    """Format text in cell"""
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(2)
        paragraph_format.line_spacing = 1.05
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor(*color)


def _set_borderless_table(table):
    """Remove all borders except bottom borders for rows (borderless style)"""
    tbl = table._element
    tblPr = tbl.tblPr
    
    # Remove all borders from table
    tblBorders = parse_xml(r'<w:tblBorders {}><w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/><w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/></w:tblBorders>'.format(nsdecls('w')))
    
    # Remove existing borders
    existing_borders = tblPr.xpath('.//w:tblBorders')
    for b in existing_borders:
        tblPr.remove(b)
    tblPr.append(tblBorders)
    
    # Set bottom border for each row (except header, which gets its own style)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            
            # Remove all existing borders
            existing_cell_borders = tcPr.xpath('.//w:tcBorders')
            for b in existing_cell_borders:
                tcPr.remove(b)
            
            # Add only bottom border for data rows (not header)
            if row_idx > 0:
                cell_borders = parse_xml(r'<w:tcBorders {}><w:bottom w:val="single" w:sz="8" w:space="0" w:color="E5E7EB"/></w:tcBorders>'.format(nsdecls('w')))
                tcPr.append(cell_borders)
            
            # Remove padding (set minimal padding)
            tcMar = parse_xml(r'<w:tcMar {}><w:top w:w="80" w:type="dxa"/><w:left w:w="80" w:type="dxa"/><w:bottom w:w="80" w:type="dxa"/><w:right w:w="80" w:type="dxa"/></w:tcMar>'.format(nsdecls('w')))
            existing_margins = tcPr.xpath('.//w:tcMar')
            for m in existing_margins:
                tcPr.remove(m)
            tcPr.append(tcMar)


def create_table():
    """Create stakeholder group table (borderless, only bottom borders)
    Table width: 10 inches (A4 landscape full width)"""
    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    table.style = None  # Remove default style
    
    # Set table width to A4 landscape (10 inches)
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = parse_xml(r'<w:tblPr></w:tblPr>')
        tbl.insert(0, tblPr)
    
    # Set table width to 10 inches (A4 landscape)
    tblW = parse_xml(r'<w:tblW {} w:w="{}" w:type="dxa"/>'.format(nsdecls('w'), 14400))  # 10 inches = 14400 twips
    existing_width = tblPr.xpath('.//w:tblW')
    for w in existing_width:
        tblPr.remove(w)
    tblPr.append(tblW)
    
    # Set column widths: 20%, 40%, 40% of 10 inches
    table.columns[0].width = Inches(2.0)  # 20% of 10 inches
    table.columns[1].width = Inches(4.0)  # 40% of 10 inches
    table.columns[2].width = Inches(4.0)  # 40% of 10 inches
    
    # Headers
    headers = ["Stakeholder Group", "Material Topics of Concern (Focus)", "Engagement Channels & Frequency"]
    header_cells = table.rows[0].cells
    
    for idx, header in enumerate(headers):
        _format_text(header_cells[idx], header, bold=True, size=11, color=(31, 41, 55))
        # Header background: light gray
        _apply_shading(header_cells[idx], "F3F4F6")
    
    # Data rows
    for data in STAKEHOLDER_DATA:
        cells = table.add_row().cells
        
        # Column 1: Stakeholder Group (bold)
        _format_text(cells[0], data["group"], bold=True, size=10, color=(31, 41, 55))
        
        # Column 2: Material Topics
        _format_text(cells[1], data["topics"], bold=False, size=10, color=(55, 65, 81))
        
        # Column 3: Engagement Channels
        _format_text(cells[2], data["channels"], bold=False, size=10, color=(55, 65, 81))
    
    # Apply borderless style (only bottom borders)
    _set_borderless_table(table)
    
    return table


if __name__ == "__main__":
    # Test: save as standalone document
    table = create_table()
    doc = Document()
    doc._element.body.append(table._element)
    
    # Add reference note
    note_para = doc.add_paragraph()
    note_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    note_run = note_para.add_run("Reference: Adapted from GRI Standards and industry-leading banking reports (BEA, HSBC).")
    note_run.font.size = Pt(9)
    note_run.font.color.rgb = RGBColor(107, 114, 128)
    note_run.font.italic = True
    note_para.paragraph_format.space_before = Pt(8)
    
    doc.save("preview_stakeholder_group.docx")
    print("Stakeholder group table generated successfully!")
    print("- preview_stakeholder_group.docx: Borderless table with bottom borders only")

