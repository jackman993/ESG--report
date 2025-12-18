from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH


HEADERS = [
    "Core Sustainability Pillars",
    "Strategic Direction",
    "Future Work Commitments",
]

ROWS = [
    {
        "pillar_title": "PILLAR ONE\nGREEN ENTERPRISE",
        "direction": "1. Green Platform Development",
        "commitments": [
            "Establish sustainable supply chain management system with strategic partners",
            "Promote green certification rates for suppliers and business partners",
            "Develop ESG-related product map for green and sustainable products",
            "Implement 'City-to-City' sustainable development plan",
            "Launch digital platform for tracking sustainability metrics",
        ],
    },
    {
        "pillar_title": "PILLAR ONE\nINTERNAL MANAGEMENT",
        "direction": "2. Internal Management Enhancement",
        "commitments": [
            "Achieve ISO certifications including ISO 14001 and ISO 50001 management systems",
            "Enhance resource management: optimize energy consumption and reduce water usage",
            "Promote waste reduction initiatives across all operational facilities",
            "Pursue green building certifications for key premises",
            "Implement energy efficiency programs with reduction targets",
        ],
    },
    {
        "pillar_title": "PILLAR TWO\nSUSTAINABLE OPERATIONS",
        "direction": "3. Promoting Long-term Value",
        "commitments": [
            "Integrate ESG criteria into employee performance evaluation",
            "Implement diversified talent development and retention plans",
            "Expand range of ESG-focused financial products and services",
            "Promote circular economy concept in business practices",
            "Establish stakeholder engagement programs for sustainability initiatives",
        ],
    },
]


def _apply_shading(cell, color_hex: str):
    shading = parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading)


def _format_cell_text(cell, text, bold=False, font_size=10, color_rgb=(255, 255, 255), align_center=False):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(3)
        for run in paragraph.runs:
            run.font.size = Pt(font_size)
            run.font.bold = bold
            if color_rgb is not None:
                run.font.color.rgb = RGBColor(*color_rgb)
        if align_center:
            paragraph.alignment = 1


def create_sustainability_strategy_table():
    """Create a sustainability strategy table for insertion into the Word report."""
    doc = Document()
    
    # Add main title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_para.add_run("Sustainability Strategy and Action Plan Overview")
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(32, 41, 63)
    title_para.paragraph_format.space_after = Pt(4)
    
    # Add subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_run = subtitle_para.add_run("Mapping key strategic focus areas to actionable commitments.")
    subtitle_run.font.size = Pt(10)  # Fixed at 10pt as specified
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    subtitle_para.paragraph_format.space_after = Pt(10)
    
    # Create table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Set column widths (参考Environment包的方式，只设置列宽，让Word自动处理table宽度)
    # 统一规则：整页table应占满整页宽度（约10.0英寸for A4 landscape）
    # Column widths: 25%, 25%, 50% of full page width
    col_widths = [Inches(2.5), Inches(2.5), Inches(5.0)]
    
    # 设置每列的宽度（Environment包的方式）
    for i, width in enumerate(col_widths):
        table.columns[i].width = width

    # Header row
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(HEADERS):
        _format_cell_text(
            hdr_cells[idx],
            header,
            bold=True,
            font_size=10,  # Fixed at 10pt as specified
            color_rgb=(255, 255, 255),
            align_center=True,
        )
        _apply_shading(hdr_cells[idx], '1E293B')

    # Body rows
    for row in ROWS:
        cells = table.add_row().cells

        _format_cell_text(
            cells[0],
            row['pillar_title'],
            bold=True,
            font_size=10,  # Fixed at 10pt as specified
            color_rgb=(255, 255, 255),
            align_center=True,
        )
        _apply_shading(cells[0], '475569')

        _format_cell_text(
            cells[1],
            row['direction'],
            bold=True,
            font_size=10,  # Fixed at 10pt as specified
            color_rgb=(30, 41, 59),
        )
        _apply_shading(cells[1], 'E2E8F0')

        commitments_text = '\n'.join(f"• {item}" for item in row['commitments'])
        _format_cell_text(
            cells[2],
            commitments_text,
            bold=False,
            font_size=10,  # Fixed at 10pt as specified
            color_rgb=(55, 65, 81),
        )

    # Ensure every cell has at least one paragraph (Word strict requirement)
    def fix_table_structure(tbl):
        for r in tbl.rows:
            for c in r.cells:
                if len(c.paragraphs) == 0:
                    c.add_paragraph("")
                # Make sure there is text container; keep it empty for no visual impact
                if c.text is None:
                    c.text = ""

    for t in doc.tables:
        fix_table_structure(t)

    # Return the complete document (with title, subtitle, and table)
    return doc


if __name__ == "__main__":
    # Quick manual test: saves a standalone document for preview
    doc = create_sustainability_strategy_table()
    doc.save("preview_sustainability_strategy.docx")
