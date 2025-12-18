"""ESG Core Pillars table generator for Word reports."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

PILLARS = [
    {
        "pillar": "Planet",
        "focus": "Climate, emissions, resource resilience",
        "initiatives": [
            "Net-zero roadmap with interim carbon budgets",
            "Renewable energy adoption and energy-efficiency upgrades",
            "Water stewardship and nature-positive initiatives",
        ],
        "color": "10B981",
    },
    {
        "pillar": "Products",
        "focus": "Circular design, packaging, sustainable sourcing",
        "initiatives": [
            "Lifecycle assessments informing eco-design decisions",
            "Closed-loop packaging and recycled content targets",
            "Supplier ESG qualification and material traceability",
        ],
        "color": "F59E0B",
    },
    {
        "pillar": "People",
        "focus": "Ethics, capability building, inclusive culture",
        "initiatives": [
            "Diversity, equity and inclusion programmes",
            "Reskilling for green and digital roles",
            "Well-being, health & safety and community engagement",
        ],
        "color": "1E3A8A",
    },
]


def _apply_shading(cell, color_hex: str):
    shading = parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading)


def _format_text(cell, text, *, bold=False, size=10, color=(31, 41, 55), align_center=False):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(2)
        paragraph_format.line_spacing = 1.05
        if align_center:
            paragraph.alignment = 1
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor(*color)


def create_table():
    doc = Document()
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    headers = ["ESG Pillar", "Strategic Focus", "Key Initiatives"]
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        _format_text(header_cells[idx], header, bold=True, size=11, color=(255, 255, 255), align_center=True)
        _apply_shading(header_cells[idx], "1F2937")

    for pillar in PILLARS:
        row_cells = table.add_row().cells
        _format_text(
            row_cells[0],
            pillar["pillar"],
            bold=True,
            size=11,
            color=(255, 255, 255),
            align_center=True,
        )
        _apply_shading(row_cells[0], pillar["color"])

        _format_text(row_cells[1], pillar["focus"], bold=False, size=10)
        _apply_shading(row_cells[1], "E2E8F0")

        initiatives_text = "\n".join(f"• {item}" for item in pillar["initiatives"])
        _format_text(row_cells[2], initiatives_text, bold=False, size=10)

    return table


if __name__ == "__main__":
    table = create_table()
    doc = Document()
    doc._element.body.append(table._element)
    doc.save("preview_esg_core_pillars.docx")
