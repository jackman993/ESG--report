from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
from datetime import datetime


def render_risk_manage(output_path="output/risk_manage_3_6.docx"):
    doc = Document()

    # A4 landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    title = doc.add_paragraph()
    r = title.add_run("Material Issues and Risk Management")
    r.bold = True; r.font.size = Pt(20); r.font.color.rgb = RGBColor(31, 78, 120)
    title.alignment = 1

    headers = [
        "Material Issues",
        "Description of Risk",
        "Type of Risk",
        "Risk Severity",
        "Risk Likelihood",
        "Mitigation Measures",
    ]

    data = [
        ["Sustainable Supply Chain", "Disruptions due to geopolitical instability or shortage of raw materials, with added exposure to compliance lapses and scope 3 data gaps.",
         "Operational", "Moderate", "High",
         "Establish supplier risk assessment, diversify sourcing, conduct regular ESG audits, and track corrective actions with pass-rate KPIs."],
        ["Climate Strategy", "Climate-related disasters may impact operations and reputation; transition risks (policy, carbon pricing) can raise costs and affect market access.",
         "Strategic", "High", "Moderate",
         "Set science-based targets, invest in renewable energy and energy storage, integrate TCFD/ISSB disclosure, and perform scenario analysis."],
        ["Waste Management", "Inefficient waste disposal could harm reputation and increase regulatory exposure, while value leakage persists without circular design.",
         "Environmental", "Low", "Moderate",
         "Expand segregation, recovery, and by-product valorization; set diversion-from-landfill targets and supplier packaging reduction KPIs."],
        ["Energy Efficiency", "Rising energy costs and limited access to clean energy increase operating volatility and climate exposure.",
         "Operational", "Moderate", "High",
         "Upgrade to high-efficiency equipment, deploy continuous monitoring, peak-shaving, and ISO 50001 processes with verified M&V."],
        ["Health & Safety", "Workplace injuries or low safety awareness; psychosocial risks may impact productivity and well-being.",
         "Operational", "Low", "Moderate",
         "Conduct targeted training, inspections, near-miss reporting, ergonomic interventions, and data-driven campaigns with TRIR KPIs."],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    try:
        table.allow_autofit = False
    except Exception:
        pass

    # Table width 9 inches (shrink by ~10%)
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr'); table._tbl.insert(0, tblPr) if table._tbl.tblPr is None else None
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), '12960'); tblW.set(qn('w:type'), 'dxa')
    for w in tblPr.xpath('.//w:tblW'):
        tblPr.remove(w)
    tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout'); tblLayout.set(qn('w:type'), 'fixed')
    for l in tblPr.xpath('.//w:tblLayout'):
        tblPr.remove(l)
    tblPr.append(tblLayout)

    # Column widths (inches) summing to 10
    col_widths = [Inches(1.8), Inches(2.8), Inches(1.0), Inches(0.8), Inches(1.0), Inches(2.6)]

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = 1
        run = p.add_run(header)
        run.bold = True; run.font.size = Pt(11); run.font.color.rgb = RGBColor(255, 255, 255)
        shading = parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="1F4E78"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        hdr_cells[i].width = col_widths[i]

    # Body
    for r_idx, row in enumerate(data):
        cells = table.add_row().cells
        for j, text in enumerate(row):
            para = cells[j].paragraphs[0]
            para.alignment = 0
            run = para.add_run(text)
            run.font.size = Pt(9); run.font.name = "Calibri"
            try:
                from docx.enum.text import WD_LINE_SPACING
                para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                para.paragraph_format.line_spacing = Pt(12)
            except Exception:
                pass
            para.paragraph_format.space_after = Pt(2)
            cells[j].width = col_widths[j]
            if r_idx % 2 == 0:
                shade = parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="F2F2F2"/>'.format(nsdecls('w')))
                cells[j]._tc.get_or_add_tcPr().append(shade)

    os.makedirs("output", exist_ok=True)
    try:
        doc.save(output_path)
        final_path = output_path
    except PermissionError:
        ts = datetime.now().strftime('%Y%m%d_%H%M%S')
        base, ext = os.path.splitext(output_path)
        fallback = f"{base}_{ts}{ext}"
        doc.save(fallback)
        final_path = fallback
        print(f"[WARN] Target file locked. Saved as: {final_path}")
    print(f"[SUCCESS] Word file successfully generated: {final_path}")


def create_table():
    """Return a Document containing only the risk management table (for embedding)."""
    doc = Document()
    # Rebuild compact table (reuse same headers/data and layout)
    headers = [
        "Material Issues","Description of Risk","Type of Risk","Risk Severity","Risk Likelihood","Mitigation Measures"
    ]
    data = [
        ["Sustainable Supply Chain","Disruptions due to geopolitical instability or shortage of raw materials, with exposure to compliance lapses and scope 3 data gaps.","Operational","Moderate","High","Supplier risk assessment, diversify sourcing, ESG audits, corrective actions tracking."],
        ["Climate Strategy","Physical and transition risks impact operations, reputation and cost base.","Strategic","High","Moderate","Science-based targets, renewables/storage, TCFD/ISSB disclosure, scenario analysis."],
        ["Waste Management","Inefficient disposal harms reputation and increases regulatory exposure.","Environmental","Low","Moderate","Segregation/recovery, by‑product valorization, diversion targets, packaging reduction."],
        ["Energy Efficiency","Rising energy costs and limited clean supply increase volatility.","Operational","Moderate","High","High‑efficiency upgrades, monitoring, peak‑shaving, ISO 50001 with verified M&V."],
        ["Health & Safety","Injuries/psychosocial risk impact productivity and wellbeing.","Operational","Low","Moderate","Training, inspections, near‑miss reporting, ergonomics, TRIR campaigns."],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"; table.autofit = False
    try:
        table.allow_autofit = False
    except Exception:
        pass

    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    if table._tbl.tblPr is None:
        table._tbl.insert(0, tblPr)
    # 9 inches = 12960 twips
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'),'12960'); tblW.set(qn('w:type'),'dxa')
    for w in tblPr.xpath('.//w:tblW'):
        tblPr.remove(w)
    tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout'); tblLayout.set(qn('w:type'),'fixed')
    for l in tblPr.xpath('.//w:tblLayout'):
        tblPr.remove(l)
    tblPr.append(tblLayout)

    col_widths = [Inches(1.8), Inches(2.8), Inches(1.0), Inches(0.8), Inches(1.0), Inches(2.6)]
    hdr_cells = table.rows[0].cells
    for i,h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]; p.alignment=1
        run = p.add_run(h); run.bold=True; run.font.size=Pt(11); run.font.color.rgb=RGBColor(255,255,255)
        shading = parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="1F4E78"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        hdr_cells[i].width = col_widths[i]

    for r in data:
        cells = table.add_row().cells
        for j,text in enumerate(r):
            p = cells[j].paragraphs[0]; p.alignment=0
            p.add_run(text).font.size = Pt(9)
            cells[j].width = col_widths[j]

    # Ensure strict cell structure to satisfy Word
    def fix_table_structure(tbl):
        for tr in tbl.rows:
            for c in tr.cells:
                if len(c.paragraphs) == 0:
                    c.add_paragraph("")
                if c.text is None:
                    c.text = ""

    for t in doc.tables:
        fix_table_structure(t)

    return doc


if __name__ == "__main__":
    render_risk_manage()
