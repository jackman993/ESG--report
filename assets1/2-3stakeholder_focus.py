"""Stakeholder focus radar chart for Word reports.
Creates a radar chart image using matplotlib and inserts it into Word."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
from io import BytesIO
import os

# Data from the HTML Chart.js configuration
RADAR_DATA = {
    "labels": ['Climate Action', 'Ethics & Compliance', 'Human Capital', 'Data Security', 'Product Safety'],
    "datasets": [
        {
            "label": 'Investors',
            "data": [90, 80, 55, 75, 60],
            "color": '#0077B6',  # Blue
            "alpha": 0.3,
        },
        {
            "label": 'Customers/Employees',
            "data": [65, 70, 95, 85, 90],
            "color": '#48CAE4',  # Cyan
            "alpha": 0.3,
        }
    ]
}


def create_radar_chart_image():
    """Create radar chart using matplotlib and return as image bytes"""
    # Number of variables
    categories = RADAR_DATA["labels"]
    N = len(categories)
    
    # Compute angle for each category
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]  # Complete the circle
    
    # Initialize the figure
    fig, ax = plt.subplots(figsize=(8, 8), subplot_kw=dict(projection='polar'))
    
    # Plot each dataset
    for dataset in RADAR_DATA["datasets"]:
        values = dataset["data"] + dataset["data"][:1]  # Complete the circle
        
        # Plot
        ax.plot(angles, values, 'o-', linewidth=2, label=dataset["label"], 
                color=dataset["color"], markersize=6)
        ax.fill(angles, values, alpha=dataset["alpha"], color=dataset["color"])
    
    # Add labels
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=9)
    
    # Set y-axis limits and labels
    ax.set_ylim(0, 100)
    ax.set_yticks([0, 20, 40, 60, 80, 100])
    ax.set_yticklabels(['0', '20', '40', '60', '80', '100'], fontsize=8)
    ax.grid(True, linestyle='--', alpha=0.5)
    
    # Add legend
    ax.legend(loc='upper right', bbox_to_anchor=(1.3, 1.1), fontsize=9)
    
    # Set title
    ax.set_title('Social: Stakeholder Focus Areas (Perception Score)', 
                 fontsize=11, fontweight='bold', pad=20, color='#00B4D8')
    
    # Save to bytes
    img_buffer = BytesIO()
    plt.tight_layout()
    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    img_buffer.seek(0)
    plt.close()
    
    return img_buffer


def create_table():
    """Create a Word document with radar chart (compact size for cell insertion)"""
    doc = Document()
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Social: Stakeholder Focus Areas (Perception Score)")
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 180, 216)  # #00B4D8
    title_para.paragraph_format.space_after = Pt(4)
    
    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Average perception score (0-100) from key stakeholder groups.")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(107, 114, 128)
    subtitle_para.paragraph_format.space_after = Pt(12)
    
    # Create radar chart image (use smaller version for cell)
    img_buffer = create_radar_chart_for_cell()
    
    # Add image to document
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(img_buffer, width=Inches(4.0))  # compact width to ensure rendering in cell
    
    return doc


def create_radar_chart_for_cell():
    """Create radar chart image for insertion into a cell (smaller size)"""
    # Number of variables
    categories = RADAR_DATA["labels"]
    N = len(categories)
    
    # Compute angle for each category
    angles = [n / float(N) * 2 * np.pi for n in range(N)]
    angles += angles[:1]  # Complete the circle
    
    # Initialize the figure (smaller for cell)
    fig, ax = plt.subplots(figsize=(5, 5), subplot_kw=dict(projection='polar'))
    
    # Plot each dataset
    for dataset in RADAR_DATA["datasets"]:
        values = dataset["data"] + dataset["data"][:1]  # Complete the circle
        
        # Plot
        ax.plot(angles, values, 'o-', linewidth=1.5, label=dataset["label"], 
                color=dataset["color"], markersize=4)
        ax.fill(angles, values, alpha=dataset["alpha"], color=dataset["color"])
    
    # Add labels
    ax.set_xticks(angles[:-1])
    ax.set_xticklabels(categories, fontsize=7)
    
    # Set y-axis limits and labels
    ax.set_ylim(0, 100)
    ax.set_yticks([0, 50, 100])
    ax.set_yticklabels(['0', '50', '100'], fontsize=7)
    ax.grid(True, linestyle='--', alpha=0.5)
    
    # Add legend
    ax.legend(loc='upper right', bbox_to_anchor=(1.25, 1.15), fontsize=7)
    
    # Set title
    ax.set_title('Stakeholder Focus Areas', fontsize=9, fontweight='bold', 
                 pad=15, color='#00B4D8')
    
    # Save to bytes
    img_buffer = BytesIO()
    plt.tight_layout()
    plt.savefig(img_buffer, format='png', dpi=120, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    img_buffer.seek(0)
    plt.close()
    
    return img_buffer


if __name__ == "__main__":
    try:
        # Test: save as standalone document
        doc = create_table()
        doc.save("preview_stakeholder_focus.docx")
        print("Stakeholder focus radar chart generated successfully!")
        print("- preview_stakeholder_focus.docx: Full-page radar chart")
        
        # Test: create chart image for cell
        img_buffer = create_radar_chart_for_cell()
        with open("preview_stakeholder_focus_chart.png", "wb") as f:
            f.write(img_buffer.getvalue())
        print("- preview_stakeholder_focus_chart.png: Radar chart image for cell insertion")
        
    except ImportError:
        print("Error: matplotlib is required. Install with: pip install matplotlib numpy")
    except Exception as e:
        print(f"Error generating radar chart: {e}")

