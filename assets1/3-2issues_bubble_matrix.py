"""ESG Materiality Assessment Matrix (Bubble Plot) for Word reports.
Creates a bubble matrix chart using matplotlib and inserts it into Word."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO

# Data from the HTML Chart.js configuration
# Format: {label, x: Financial Impact (0-100), y: Stakeholder Importance (0-100), r: Bubble Size (Issue Urgency), color}
BUBBLE_DATA = [
    # Quadrant 1: High Materiality (Strategic Focus)
    {'label': 'Climate Transition Strategy', 'x': 90, 'y': 85, 'r': 15, 'color': '#0077B6'},
    {'label': 'Supply Chain Human Rights', 'x': 65, 'y': 95, 'r': 12, 'color': '#00B4D8'},
    {'label': 'AI Ethics & Data Privacy', 'x': 80, 'y': 70, 'r': 10, 'color': '#48CAE4'},
    {'label': 'Talent Retention & DEI', 'x': 70, 'y': 80, 'r': 10, 'color': '#90E0EF'},
    
    # Quadrant 2: Financial Focus (High Impact, Low Stakeholder)
    {'label': 'Cyber Security Governance', 'x': 95, 'y': 40, 'r': 14, 'color': '#0077B6'},
    {'label': 'Sustainable Finance/Tax', 'x': 85, 'y': 35, 'r': 9, 'color': '#00B4D8'},
    
    # Quadrant 3: Stakeholder Focus (Low Impact, High Stakeholder)
    {'label': 'Community Engagement', 'x': 20, 'y': 80, 'r': 7, 'color': '#90E0EF'},
    
    # Quadrant 4: Lower Materiality (Monitoring)
    {'label': 'Waste & Circularity', 'x': 45, 'y': 55, 'r': 7, 'color': '#CAF0F8'},
    {'label': 'Responsible Marketing', 'x': 30, 'y': 45, 'r': 5, 'color': '#CAF0F8'},
    {'label': 'Product Packaging', 'x': 55, 'y': 20, 'r': 6, 'color': '#CAF0F8'},
]


def hex_to_rgb(hex_color):
    """Convert hex color to RGB tuple (0-1 range for matplotlib)"""
    hex_color = hex_color.lstrip('#')
    rgb = tuple(int(hex_color[i:i+2], 16) / 255.0 for i in (0, 2, 4))
    return rgb


def create_bubble_matrix_image():
    """Create bubble matrix chart using matplotlib and return as image bytes"""
    # Initialize the figure (reduced by 40%: 10*0.6=6, 8*0.6=4.8)
    fig, ax = plt.subplots(figsize=(6, 4.8))
    
    # Extract data
    x_values = [item['x'] for item in BUBBLE_DATA]
    y_values = [item['y'] for item in BUBBLE_DATA]
    sizes = [item['r'] * 200 for item in BUBBLE_DATA]  # Increase bubble size significantly
    colors = [hex_to_rgb(item['color']) for item in BUBBLE_DATA]
    labels = [item['label'] for item in BUBBLE_DATA]
    
    # Create scatter plot (bubbles) - increase alpha and make sure bubbles are visible
    scatter = ax.scatter(x_values, y_values, s=sizes, c=colors, 
                        alpha=0.7, edgecolors='white', linewidths=2)
    
    # Add quadrant divider lines at 50, 50
    ax.axhline(y=50, color='gray', linestyle='--', linewidth=1, alpha=0.5)
    ax.axvline(x=50, color='gray', linestyle='--', linewidth=1, alpha=0.5)
    
    # Add quadrant labels (smaller for 40% reduction)
    ax.text(25, 75, 'Stakeholder\nFocus', ha='center', va='center', 
            fontsize=7, color='gray', style='italic', alpha=0.7)
    ax.text(75, 75, 'Strategic\nFocus', ha='center', va='center', 
            fontsize=7, color='gray', style='italic', alpha=0.7, fontweight='bold')
    ax.text(25, 25, 'Monitoring', ha='center', va='center', 
            fontsize=7, color='gray', style='italic', alpha=0.7)
    ax.text(75, 25, 'Financial\nFocus', ha='center', va='center', 
            fontsize=7, color='gray', style='italic', alpha=0.7)
    
    # Add labels to bubbles (simpler labels, smaller font for compact display)
    for item in BUBBLE_DATA:
        # Use first 2-3 words for label
        words = item['label'].split()
        if len(words) > 2:
            label = ' '.join(words[:2])  # First 2 words
        else:
            label = item['label']
        
        ax.text(item['x'], item['y'], label, ha='center', va='center', 
                fontsize=6, fontweight='bold', color='white',
                bbox=dict(boxstyle='round,pad=0.2', facecolor='black', alpha=0.6, edgecolor='none'))
    
    # Set axis limits
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    
    # Set axis labels (smaller for 40% reduction)
    ax.set_xlabel('Financial Impact (Low → High)', fontsize=9, fontweight='bold')
    ax.set_ylabel('Stakeholder Importance (Low → High)', fontsize=9, fontweight='bold')
    
    # Set axis ticks (fewer ticks for compact display)
    ax.set_xticks([0, 50, 100])
    ax.set_xticklabels(['0', '50', '100'], fontsize=8)
    ax.set_yticks([0, 50, 100])
    ax.set_yticklabels(['0', '50', '100'], fontsize=8)
    
    # Grid
    ax.grid(True, linestyle='--', alpha=0.3, linewidth=0.5)
    ax.set_axisbelow(True)
    
    # Remove top and right spines
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    # Set title (smaller for 40% reduction)
    ax.set_title('ESG Materiality Assessment Matrix (Bubble Plot)', 
                 fontsize=10, fontweight='bold', pad=10, color='#0077B6')
    
    # Tight layout
    plt.tight_layout()
    
    # Save to bytes
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    img_buffer.seek(0)
    plt.close()
    
    return img_buffer


def create_table():
    """Create a Word document with bubble matrix chart"""
    doc = Document()
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("ESG Materiality Assessment Matrix (Bubble Plot)")
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 119, 182)  # #0077B6
    title_para.paragraph_format.space_after = Pt(4)
    
    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Evaluating issues based on Stakeholder Importance (Y-axis) vs. Financial Impact (X-axis).")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(107, 114, 128)
    subtitle_para.paragraph_format.space_after = Pt(12)
    
    # Create bubble matrix chart image
    img_buffer = create_bubble_matrix_image()
    
    # Add image to document (reduced by 40%: 8 * 0.6 = 4.8 inches)
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(img_buffer, width=Inches(4.8))  # 40% smaller: 4.8 inches wide
    
    return doc


def create_bubble_matrix_for_cell():
    """Create bubble matrix image for insertion into a cell (smaller size)"""
    # Initialize the figure (smaller for cell)
    fig, ax = plt.subplots(figsize=(6, 6))
    
    # Extract data
    x_values = [item['x'] for item in BUBBLE_DATA]
    y_values = [item['y'] for item in BUBBLE_DATA]
    sizes = [item['r'] * 15 for item in BUBBLE_DATA]  # Smaller scale for compact display
    colors = [hex_to_rgb(item['color']) for item in BUBBLE_DATA]
    labels = [item['label'] for item in BUBBLE_DATA]
    
    # Create scatter plot (bubbles)
    scatter = ax.scatter(x_values, y_values, s=sizes, c=colors, 
                        alpha=0.6, edgecolors='white', linewidths=1)
    
    # Add quadrant divider lines at 50, 50
    ax.axhline(y=50, color='gray', linestyle='--', linewidth=0.8, alpha=0.5)
    ax.axvline(x=50, color='gray', linestyle='--', linewidth=0.8, alpha=0.5)
    
    # Add quadrant labels (smaller)
    ax.text(25, 75, 'Stakeholder\nFocus', ha='center', va='center', 
            fontsize=8, color='gray', style='italic', alpha=0.7)
    ax.text(75, 75, 'Strategic\nFocus', ha='center', va='center', 
            fontsize=8, color='gray', style='italic', alpha=0.7, fontweight='bold')
    ax.text(25, 25, 'Monitoring', ha='center', va='center', 
            fontsize=8, color='gray', style='italic', alpha=0.7)
    ax.text(75, 25, 'Financial\nFocus', ha='center', va='center', 
            fontsize=8, color='gray', style='italic', alpha=0.7)
    
    # Add labels to bubbles (smaller font, shorter labels)
    for item in BUBBLE_DATA:
        label = item['label']
        # Use first few words for compact display
        words = label.split()
        if len(words) > 2:
            label = ' '.join(words[:2])  # Use first 2 words only
        
        ax.text(item['x'], item['y'], label, ha='center', va='center', 
                fontsize=6, fontweight='bold', color='white',
                bbox=dict(boxstyle='round,pad=0.2', facecolor='black', alpha=0.5, edgecolor='none'))
    
    # Set axis limits
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    
    # Set axis labels (smaller)
    ax.set_xlabel('Financial Impact', fontsize=9, fontweight='bold')
    ax.set_ylabel('Stakeholder Importance', fontsize=9, fontweight='bold')
    
    # Set axis ticks (fewer ticks)
    ax.set_xticks([0, 50, 100])
    ax.set_xticklabels(['0', '50', '100'], fontsize=7)
    ax.set_yticks([0, 50, 100])
    ax.set_yticklabels(['0', '50', '100'], fontsize=7)
    
    # Grid
    ax.grid(True, linestyle='--', alpha=0.3, linewidth=0.5)
    ax.set_axisbelow(True)
    
    # Remove top and right spines
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    
    # Set title (smaller)
    ax.set_title('Materiality Assessment Matrix', 
                 fontsize=10, fontweight='bold', pad=10, color='#0077B6')
    
    # Tight layout
    plt.tight_layout()
    
    # Save to bytes
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', dpi=120, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    img_buffer.seek(0)
    plt.close()
    
    return img_buffer


if __name__ == "__main__":
    import time
    try:
        # Test: save as standalone document
        doc = create_table()
        timestamp = int(time.time())
        filename1 = f"preview_bubble_matrix_{timestamp}.docx"
        doc.save(filename1)
        print("ESG Materiality Assessment Matrix generated successfully!")
        print(f"- {filename1}: Full-page bubble matrix chart (40% smaller, larger bubbles)")
        
        # Test: create chart image for cell
        img_buffer = create_bubble_matrix_for_cell()
        filename2 = f"preview_bubble_matrix_chart_{timestamp}.png"
        with open(filename2, "wb") as f:
            f.write(img_buffer.getvalue())
        print(f"- {filename2}: Bubble matrix image for cell insertion")
        
    except ImportError:
        print("Error: matplotlib is required. Install with: pip install matplotlib numpy")
    except Exception as e:
        print(f"Error generating bubble matrix chart: {e}")

