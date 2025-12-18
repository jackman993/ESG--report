"""Strategic material issues bar chart for Word reports.
Creates a horizontal bar chart using matplotlib and inserts it into Word."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO

# Data from the HTML Chart.js configuration
BAR_DATA = {
    "issues": [
        'Climate Transition Strategy', 'Talent Retention', 'AI Ethics & Governance', 
        'Supply Chain Human Rights', 'Water Stewardship', 'Data Privacy', 
        'Circular Economy', 'Board Diversity'
    ],
    "scores": [95, 88, 85, 80, 72, 68, 60, 55],
    "colors": ['#0077B6', '#00B4D8', '#48CAE4', '#90E0EF', '#CAF0F8', 
               '#0077B6', '#00B4D8', '#48CAE4'],  # Gradient blue colors
}


def create_bar_chart_image():
    """Create horizontal bar chart using matplotlib and return as image bytes"""
    # Initialize the figure
    fig, ax = plt.subplots(figsize=(10, 6))
    
    # Reverse order for top-to-bottom display (highest first)
    issues = BAR_DATA["issues"][::-1]
    scores = BAR_DATA["scores"][::-1]
    colors = BAR_DATA["colors"][::-1]
    
    # Create horizontal bar chart
    y_pos = np.arange(len(issues))
    bars = ax.barh(y_pos, scores, color=colors, alpha=0.8, edgecolor='white', linewidth=1)
    
    # Set y-axis labels
    ax.set_yticks(y_pos)
    ax.set_yticklabels(issues, fontsize=9)
    
    # Set x-axis limits and labels
    ax.set_xlim(0, 100)
    ax.set_xlabel('Prioritization Score (0-100)', fontsize=10, fontweight='bold')
    ax.set_xticks([0, 20, 40, 60, 80, 100])
    ax.set_xticklabels(['0', '20', '40', '60', '80', '100'], fontsize=8)
    
    # Add value labels on bars
    for i, (bar, score) in enumerate(zip(bars, scores)):
        width = bar.get_width()
        ax.text(width + 1, bar.get_y() + bar.get_height()/2, 
                f'{score}', ha='left', va='center', fontsize=9, fontweight='bold')
    
    # Remove top and right spines
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.spines['bottom'].set_visible(True)
    
    # Remove y-axis grid
    ax.grid(axis='x', linestyle='--', alpha=0.3, linewidth=0.5)
    ax.set_axisbelow(True)
    
    # Set title
    ax.set_title('Strategic Material Issues Ranking (Prioritization Score)', 
                 fontsize=12, fontweight='bold', pad=15, color='#48CAE4')
    
    # Tight layout
    plt.tight_layout()
    
    # Save to bytes
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', dpi=150, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    img_buffer.seek(0)
    plt.close()
    
    return img_buffer


def create_table():
    """Create a Word document with bar chart"""
    doc = Document()
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("Strategic Material Issues Ranking (Prioritization Score)")
    title_run.font.size = Pt(14)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(72, 202, 228)  # #48CAE4
    title_para.paragraph_format.space_after = Pt(4)
    
    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Ranking of the top 8 issues based on combined financial and impact materiality scores.")
    subtitle_run.font.size = Pt(10)
    subtitle_run.font.color.rgb = RGBColor(107, 114, 128)
    subtitle_para.paragraph_format.space_after = Pt(12)
    
    # Create bar chart image
    img_buffer = create_bar_chart_image()
    
    # Add image to document
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(img_buffer, width=Inches(8))  # 8 inches wide for full-page
    
    return doc


def create_bar_chart_for_cell():
    """Create bar chart image for insertion into a cell (smaller size)"""
    # Initialize the figure (smaller for cell)
    fig, ax = plt.subplots(figsize=(6, 4.5))
    
    # Reverse order for top-to-bottom display (highest first)
    issues = BAR_DATA["issues"][::-1]
    scores = BAR_DATA["scores"][::-1]
    colors = BAR_DATA["colors"][::-1]
    
    # Create horizontal bar chart
    y_pos = np.arange(len(issues))
    bars = ax.barh(y_pos, scores, color=colors, alpha=0.8, edgecolor='white', linewidth=0.8)
    
    # Set y-axis labels (smaller font for compact display)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(issues, fontsize=7)
    
    # Set x-axis limits and labels
    ax.set_xlim(0, 100)
    ax.set_xlabel('Score', fontsize=8, fontweight='bold')
    ax.set_xticks([0, 50, 100])
    ax.set_xticklabels(['0', '50', '100'], fontsize=7)
    
    # Add value labels on bars (compact)
    for i, (bar, score) in enumerate(zip(bars, scores)):
        width = bar.get_width()
        ax.text(width + 1, bar.get_y() + bar.get_height()/2, 
                f'{score}', ha='left', va='center', fontsize=7, fontweight='bold')
    
    # Remove top and right spines
    ax.spines['top'].set_visible(False)
    ax.spines['right'].set_visible(False)
    ax.spines['left'].set_visible(False)
    ax.spines['bottom'].set_visible(True)
    
    # Remove y-axis grid
    ax.grid(axis='x', linestyle='--', alpha=0.3, linewidth=0.5)
    ax.set_axisbelow(True)
    
    # Set title (smaller)
    ax.set_title('Strategic Material Issues Ranking', 
                 fontsize=9, fontweight='bold', pad=10, color='#48CAE4')
    
    # Tight layout
    plt.tight_layout()
    
    # Save to bytes
    img_buffer = BytesIO()
    plt.savefig(img_buffer, format='png', dpi=120, bbox_inches='tight', 
                facecolor='white', edgecolor='none')
    img_buffer.seek(0)
    plt.close()
    
    return img_buffer


if __name__ == "__main__":
    try:
        # Test: save as standalone document
        doc = create_table()
        doc.save("preview_big_issues_bar.docx")
        print("Strategic material issues bar chart generated successfully!")
        print("- preview_big_issues_bar.docx: Full-page horizontal bar chart")
        
        # Test: create chart image for cell
        img_buffer = create_bar_chart_for_cell()
        with open("preview_big_issues_bar_chart.png", "wb") as f:
            f.write(img_buffer.getvalue())
        print("- preview_big_issues_bar_chart.png: Bar chart image for cell insertion")
        
    except ImportError:
        print("Error: matplotlib is required. Install with: pip install matplotlib numpy")
    except Exception as e:
        print(f"Error generating bar chart: {e}")

