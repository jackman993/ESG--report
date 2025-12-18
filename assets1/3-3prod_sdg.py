# ==========================================
# File: prod_sdg.py
# Title: Products and Services in Response to Sustainability Goals
# Version: Stable Release (No SDG Icons)
# ==========================================

from docx import Document
from docx.shared import Pt, RGBColor, Mm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from pathlib import Path
from datetime import datetime

def render_prod_sdg(output_path="output/prod_sdg.docx"):
    """
    Generate a clean Word table describing key products and services
    in response to sustainability goals. This version uses pure text only.
    """

    # Initialize document
    doc = Document()

    # ---------- Page setup: A4 Landscape with balanced margins ----------
    section = doc.sections[0]
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.left_margin = Mm(20)
    section.right_margin = Mm(20)
    section.top_margin = Mm(12)
    section.bottom_margin = Mm(12)

    # ---------- Base font ----------
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)

    # ---------- Title ----------
    title = doc.add_heading("Products and Services in Response to Sustainability Goals", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add banner-like shading to the title paragraph to resemble report header
    title_p = title._p
    title_p.get_or_add_pPr().append(
        parse_xml(r'<w:shd {} w:val="clear" w:fill="E9EEF6"/>'.format(nsdecls('w')))
    )

    # ---------- Introductory Paragraph ----------
    intro = doc.add_paragraph(
        "This table presents our core product and service portfolio aligned to material ESG topics, "
        "demonstrating how we translate strategy into measurable outcomes across energy, resources, "
        "and human capital. As an ESG-driven organization, we prioritize lifecycle thinking, climate "
        "resilience, circularity, and inclusive growth. Each solution is designed to reduce emissions, "
        "improve operational efficiency, and strengthen stakeholder trust through transparent governance "
        "and data integrity. We also emphasize standards alignment (e.g., GHG Protocol, ISO 50001/14001) "
        "and continuous improvement, ensuring that performance can be tracked with auditable KPIs and "
        "benchmarked against peer best practices."
    )
    intro_format = intro.paragraph_format
    intro_format.space_after = Pt(6)
    intro_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    intro_format.line_spacing = Pt(12)
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # ---------- Table Data ----------
    headers = ["Product / Service Category", "Description", "Applied Technology"]
    data = [
        [
            "Energy Management System",
            "End-to-end energy performance management that identifies abatement opportunities, optimizes load profiles, "
            "and enables a credible renewable energy roadmap (on-site PV, PPAs, RECs). The solution integrates baseline "
            "establishment, target setting, and M&V processes consistent with ISO 50001 while supporting Scope 2 reduction "
            "strategies (market- and location-based accounting). It drives cost savings, peak shaving, and emissions "
            "reductions with audit-ready evidence for assurance.",
            "IoT metering, BMS/SCADA integrations, edge analytics, cloud data lake, automated anomaly detection, digital twins, "
            "and dashboarding with API export to ESG reporting systems."
        ],
        [
            "Smart Manufacturing Service",
            "Data-driven production optimization to reduce scrap, water and energy intensity, and unplanned downtime. "
            "Applies lean principles and real-time quality analytics to minimize material loss and improve OEE. "
            "Supports circularity through by-product valorization and design-for-recyclability guidance, while embedding "
            "risk controls for occupational health and safety and ethical sourcing.",
            "Machine learning for predictive maintenance, process mining, MES/ERP integrations, computer vision QC, "
            "and digital work instructions connected to a compliance knowledge base."
        ],
        [
            "Digital Learning Platform",
            "Enterprise learning ecosystem that scales ESG literacy, climate risk awareness, and green skills reskilling/upskilling. "
            "Enables equitable access to training, supports talent mobility, and nurtures a culture of responsibility and "
            "innovation. Includes curated tracks on decarbonization, TCFD/ISSB disclosure readiness, and supplier engagement.",
            "Cloud LMS, adaptive learning, micro-credentialing, analytics for learning outcomes, content authoring tools, "
            "and secure SSO integration for cross-organization collaboration."
        ],
        [
            "Employee Wellness Program",
            "Holistic health and well-being framework covering physical, mental, and social dimensions to enhance productivity, "
            "retention, and psychological safety. Provides confidential support, early risk detection, ergonomic interventions, "
            "and equitable benefits. Program KPIs align with human capital reporting and DEI goals while safeguarding privacy.",
            "Wearables integration, privacy-preserving analytics, telehealth platforms, engagement apps, and secure HRIS linkage "
            "for anonymized trend reporting and impact measurement."
        ],
        [
            "Green Supply Chain",
            "Supplier enablement and performance management program that advances low-carbon procurement, waste minimization, "
            "and responsible sourcing. Incorporates Code of Conduct, onboarding, audits, and capacity building. Supports product "
            "LCAs, packaging reduction, and inbound logistics optimization to lower Scope 3 emissions with traceability.",
            "Supplier scorecards, traceability platforms, e-procurement with ESG criteria, route optimization, lifecycle databases, "
            "and automated calculation engines aligned to GHG Protocol Category mappings."
        ],
        [
            "Environmental Sensing System",
            "Continuous monitoring of ambient air quality, temperature, humidity, and noise to inform site-level risk controls and "
            "community impact management. Enables early warning for extreme weather and heat stress, supports biodiversity "
            "initiatives, and feeds climate risk assessments in line with TCFD/ISSB.",
            "Low-power IoT sensors, edge processing, satellite data fusion, geospatial analytics, event-driven alerting, and secure "
            "data pipelines to risk dashboards and incident response workflows."
        ],
    ]

    # ---------- Create Table ----------
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    # Fix layout to page width by disabling Word's autofit and setting explicit widths
    table.autofit = False
    try:
        table.allow_autofit = False
    except Exception:
        pass
    # Center the table and set a strict preferred width for A4 landscape full page (10 inches)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Add preferred table width in a version-safe way
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    # Set table width to A4 landscape full width: 10 inches = 14400 twips
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '14400')  # 10 inches = 14400 twips
    tblW.set(qn('w:type'), 'dxa')
    # Remove existing width if any
    existing_width = tblPr.xpath('.//w:tblW')
    for w in existing_width:
        tblPr.remove(w)
    tblPr.append(tblW)
    # Fixed layout so column widths are honored
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    # Remove existing layout if any
    existing_layout = tblPr.xpath('.//w:tblLayout')
    for l in existing_layout:
        tblPr.remove(l)
    tblPr.append(tblLayout)

    # Header styling
    hdr_cells = table.rows[0].cells
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        for paragraph in hdr_cells[i].paragraphs:
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(255, 255, 255)
            pf = paragraph.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.line_spacing = Pt(12)
        shading = parse_xml(r'<w:shd {} w:val="clear" w:fill="1E3A5F"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Set explicit column widths to fit A4 landscape full width (10 inches)
    # Column widths: 20% (2"), 50% (5"), 30% (3") of 10 inches
    col_widths = [Inches(2.0), Inches(5.0), Inches(3.0)]
    # Define table grid to enforce exact column widths (in twips)
    # 1 inch = 1440 twips
    col_widths_twips = [2880, 7200, 4320]  # 2", 5", 3" in twips
    
    tblGrid = tbl.tblGrid
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.append(tblGrid)
    else:
        # clear existing gridCol if any
        for child in list(tblGrid):
            tblGrid.remove(child)
    
    for w_twips in col_widths_twips:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(w_twips))
        tblGrid.append(gridCol)
    
    # Apply column widths to table
    for col_idx, w in enumerate(col_widths):
        try:
            table.columns[col_idx].width = w
        except Exception:
            pass
        for row in table.rows:
            row.cells[col_idx].width = w

    # Add rows
    for row_data in data:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            row_cells[i].text = text
            for paragraph in row_cells[i].paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)  # slightly smaller to keep width on one page
                pf = paragraph.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(1)
                pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                pf.line_spacing = Pt(12)
            row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        # Light gray background for the first column in body rows
        first_col_shading = parse_xml(r'<w:shd {} w:val="clear" w:fill="F2F2F2"/>'.format(nsdecls('w')))
        row_cells[0]._tc.get_or_add_tcPr().append(first_col_shading)

    # Prevent rows from splitting across pages for better readability
    for row in table.rows:
        row._tr.get_or_add_trPr().append(parse_xml(r'<w:cantSplit xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

    # ---------- Adjust spacing ----------
    doc.add_paragraph("\n")

    # ---------- Save File (robust, handle file-in-use) ----------
    out_path = Path(output_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(out_path))
        final_path = out_path
    except PermissionError:
        # If the file is open/locked, save with a timestamp suffix
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        fallback = out_path.with_name(f"{out_path.stem}_{timestamp}{out_path.suffix}")
        doc.save(str(fallback))
        final_path = fallback
        print(f"[WARN] Target file locked. Saved as: {final_path}")
    print(f"[SUCCESS] Word file successfully generated: {final_path}")


def create_table():
    """Build and return a Document containing the Products & Services table (for embedding)."""
    # Rebuild a compact version focused on the table only (no title/intro)
    doc = Document()

    headers = ["Product / Service Category", "Description", "Applied Technology"]
    data = [
        [
            "Energy Management System",
            "End-to-end energy performance management that identifies abatement opportunities, optimizes load profiles, and enables a credible renewable energy roadmap (on-site PV, PPAs, RECs). The solution integrates baseline establishment, target setting, and M&V processes consistent with ISO 50001.",
            "IoT metering, BMS/SCADA integrations, edge analytics, cloud data lake, anomaly detection, digital twins, dashboards.",
        ],
        [
            "Smart Manufacturing Service",
            "Data-driven production optimization to reduce scrap, intensity, and downtime; embed OH&S controls and ethical sourcing.",
            "Predictive maintenance ML, process mining, MES/ERP integrations, computer vision QC, digital work instructions.",
        ],
        [
            "Digital Learning Platform",
            "Enterprise learning to scale ESG literacy, climate risk awareness, and green skills with inclusive access.",
            "Cloud LMS, adaptive learning, micro-credentials, outcomes analytics, authoring tools, SSO integration.",
        ],
        [
            "Employee Wellness Program",
            "Holistic health and well-being with confidential support, early risk detection, and equitable benefits.",
            "Wearables, privacy-preserving analytics, telehealth, engagement apps, secure HRIS linkage.",
        ],
        [
            "Green Supply Chain",
            "Supplier enablement for low-carbon procurement, waste minimization, responsible sourcing, onboarding and audits.",
            "Scorecards, traceability platforms, e-procurement with ESG criteria, route optimization, lifecycle databases.",
        ],
        [
            "Environmental Sensing System",
            "Continuous monitoring to inform risk controls and community impact; supports biodiversity and feeds TCFD/ISSB.",
            "Low-power IoT sensors, edge processing, satellite fusion, geospatial analytics, alerting, secure pipelines.",
        ],
    ]

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    try:
        table.allow_autofit = False
    except Exception:
        pass

    # Set table width to 10 inches and fixed layout
    tbl = table._tbl
    tblPr = tbl.tblPr or OxmlElement('w:tblPr')
    if table._tbl.tblPr is None:
        table._tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW'); tblW.set(qn('w:w'), '14400'); tblW.set(qn('w:type'), 'dxa')
    for w in tblPr.xpath('.//w:tblW'):
        tblPr.remove(w)
    tblPr.append(tblW)
    tblLayout = OxmlElement('w:tblLayout'); tblLayout.set(qn('w:type'), 'fixed')
    for l in tblPr.xpath('.//w:tblLayout'):
        tblPr.remove(l)
    tblPr.append(tblLayout)

    # Column widths 2" / 5" / 3"
    col_widths = [Inches(2.0), Inches(5.0), Inches(3.0)]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            if p.runs:
                r = p.runs[0]
                r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(255, 255, 255)
            p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(0)
        shading = parse_xml(r'<w:shd {} w:val="clear" w:fill="1E3A5F"/>'.format(nsdecls('w')))
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
        hdr_cells[i].width = col_widths[i]

    for row in data:
        cells = table.add_row().cells
        for j, text in enumerate(row):
            cells[j].text = text
            for p in cells[j].paragraphs:
                p.paragraph_format.space_after = Pt(1)
                for r in p.runs:
                    r.font.size = Pt(10)
            cells[j].width = col_widths[j]

    # Ensure strict cell structure to satisfy Word
    def fix_table_structure(tbl):
        for r in tbl.rows:
            for c in r.cells:
                if len(c.paragraphs) == 0:
                    c.add_paragraph("")
                if c.text is None:
                    c.text = ""

    for t in doc.tables:
        fix_table_structure(t)

    return doc


# ---------- Main Test (for standalone run) ----------
if __name__ == "__main__":
    render_prod_sdg()
