"""ESG Roadmap timeline flowchart generator for Word reports.
Recreates the original HTML timeline design with vertical timeline, circular markers, rounded boxes, and shadows."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROADMAP_EVENTS = [
    {
        "year": "2022",
        "title": "Foundational Commitment",
        "subtitle": "Establishment & Governance",
        "items": [
            "**Governance:** Established Board-level ESG Committee and internal Steering Committee.",
            "**Commitment:** Announced Net Zero target for operational and financed emissions.",
            "**Measurement:** Initiated initial Scope 1 & 2 GHG emissions data collection and baseline setting.",
        ],
        "position": "right",
    },
    {
        "year": "2023",
        "title": "Risk and Strategy Integration",
        "subtitle": "Policy Development & Risk Mapping",
        "items": [
            "**Risk:** Integrated climate risk into Enterprise Risk Management (ERM) framework.",
            "**Social:** Mandated comprehensive Ethics and Anti-Corruption training for all staff.",
            "**Finance:** Developed Green and Sustainable Finance (GSF) Framework and sector policies.",
            "**Board:** Achieved 100% director participation in mandatory climate-related training.",
        ],
        "position": "left",
    },
    {
        "year": "2024",
        "title": "Operational Decarbonization & Due Diligence",
        "subtitle": "Execution & New Metrics",
        "items": [
            "**Operations:** Launch of Net Zero execution plans, focusing on efficiency and renewable energy adoption.",
            "**Supply Chain:** Implementation of Sustainable Procurement Policy and ESG supplier screening tools.",
            "**Governance:** Linking executive compensation to key long-term ESG performance indicators.",
            "**Data:** Completion of portfolio-wide financed emissions (Scope 3) baseline measurement.",
        ],
        "position": "right",
    },
    {
        "year": "2025",
        "title": "Forward-Looking Compliance",
        "subtitle": "Refinement & Global Standards",
        "items": [
            "**Compliance:** Full alignment with emerging IFRS S1/S2 (ISSB) financial disclosure requirements.",
            "**Engagement:** Formalized customer transition planning engagement strategy for high-impact clients.",
            "**Technology:** Integration of AI Ethics and Data Governance policy into all new product development cycles.",
            "**Review:** Conduct independent, external assurance of all major ESG disclosures and targets.",
        ],
        "position": "left",
    },
]


def _apply_shading(cell, color_hex: str):
    """Apply background shading to cell"""
    shading = parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading)


def _add_text_with_bold_keywords(paragraph, text):
    """Add text with bold keywords (words wrapped in **)"""
    paragraph.clear()
    parts = text.split("**")
    for i, part in enumerate(parts):
        if part:
            run = paragraph.add_run(part)
            # Bold keywords (odd indices after split on **)
            if i % 2 == 1:
                run.font.bold = True
                run.font.color.rgb = RGBColor(30, 41, 59)


def create_roadmap():
    """Create full roadmap document with timeline flowchart layout matching original HTML design"""
    doc = Document()
    
    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("ESG Strategy Execution Roadmap (2022 - 2025)")
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(30, 41, 59)
    title_para.paragraph_format.space_after = Pt(6)
    
    # Subtitle
    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle_para.add_run("Key Milestones for Achieving Sustainable Growth & Net Zero Targets")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
    subtitle_para.paragraph_format.space_after = Pt(20)
    
    # Create timeline layout using table
    # Each event gets a row: [left content | timeline | right content] or just [timeline | content]
    for event_idx, event in enumerate(ROADMAP_EVENTS):
        # Spacing before event (except first)
        if event_idx > 0:
            doc.add_paragraph().paragraph_format.space_before = Pt(10)
        
        # Create table row for timeline layout
        # For right-aligned: [timeline (narrow) | content (wide)]
        # For left-aligned: [content (wide) | timeline (narrow)]
        if event["position"] == "right":
            table = doc.add_table(rows=1, cols=2)
            timeline_cell = table.columns[0].cells[0]
            content_cell = table.columns[1].cells[0]
        else:
            table = doc.add_table(rows=1, cols=2)
            content_cell = table.columns[0].cells[0]
            timeline_cell = table.columns[1].cells[0]
        
        table.style = None  # Remove default style
        
        # Set cell widths
        timeline_cell.width = Inches(0.5)
        content_cell.width = Inches(7.5)
        
        # Remove all borders and padding initially
        for cell in table.rows[0].cells:
            tcPr = cell._tc.get_or_add_tcPr()
            # Remove existing borders
            for tag in ['top', 'left', 'bottom', 'right']:
                borders = tcPr.xpath(f'.//w:{tag}')
                for b in borders:
                    tcPr.remove(b)
            
            # Remove padding
            tcMar = parse_xml(r'<w:tcMar {}><w:top w:w="0" w:type="dxa"/><w:left w:w="0" w:type="dxa"/><w:bottom w:w="0" w:type="dxa"/><w:right w:w="0" w:type="dxa"/></w:tcMar>'.format(nsdecls('w')))
            existing_margins = tcPr.xpath('.//w:tcMar')
            for m in existing_margins:
                tcPr.remove(m)
            tcPr.append(tcMar)
        
        # Timeline cell: Add vertical line
        tcPr = timeline_cell._tc.get_or_add_tcPr()
        if event["position"] == "right":
            # Vertical line on right side (left side of content)
            right_border = parse_xml(r'<w:tcBorders {}><w:right w:val="single" w:sz="24" w:space="0" w:color="374151"/></w:tcBorders>'.format(nsdecls('w')))
            tcPr.append(right_border)
        else:
            # Vertical line on left side (right side of content)
            left_border = parse_xml(r'<w:tcBorders {}><w:left w:val="single" w:sz="24" w:space="0" w:color="374151"/></w:tcBorders>'.format(nsdecls('w')))
            tcPr.append(left_border)
        
        # Add circular marker in timeline cell
        timeline_para = timeline_cell.paragraphs[0]
        timeline_para.clear()
        timeline_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timeline_para.paragraph_format.space_before = Pt(8)
        timeline_para.paragraph_format.space_after = Pt(0)
        
        # Create circular marker (green circle)
        marker_run = timeline_para.add_run("●")
        marker_run.font.size = Pt(12)
        marker_run.font.color.rgb = RGBColor(16, 185, 129)  # Emerald green #10B981
        
        # Content cell: Add content box with background, border, and rounded corners effect
        # Apply light gray background (bg-slate-50)
        _apply_shading(content_cell, "F8FAFC")
        
        # Add border (left border for right-aligned, right border for left-aligned)
        tcPr = content_cell._tc.get_or_add_tcPr()
        if event["position"] == "right":
            # Left border (thick, dark)
            border = parse_xml(r'<w:tcBorders {}><w:left w:val="single" w:sz="480" w:space="0" w:color="1F2937"/></w:tcBorders>'.format(nsdecls('w')))
        else:
            # Right border (thick, dark)
            border = parse_xml(r'<w:tcBorders {}><w:right w:val="single" w:sz="480" w:space="0" w:color="1F2937"/></w:tcBorders>'.format(nsdecls('w')))
        tcPr.append(border)
        
        # Add padding to content cell
        tcMar = parse_xml(r'<w:tcMar {}><w:top w:w="144" w:type="dxa"/><w:left w:w="144" w:type="dxa"/><w:bottom w:w="144" w:type="dxa"/><w:right w:w="144" w:type="dxa"/></w:tcMar>'.format(nsdecls('w')))
        tcPr.append(tcMar)
        
        # Add content to content cell
        content_para = content_cell.paragraphs[0]
        content_para.paragraph_format.space_before = Pt(0)
        content_para.paragraph_format.space_after = Pt(0)
        
        # Year and title
        year_title_para = content_cell.add_paragraph()
        year_title_text = f"{event['year']}: {event['title']}"
        year_title_run = year_title_para.add_run(year_title_text)
        year_title_run.font.size = Pt(14)
        year_title_run.font.bold = True
        year_title_run.font.color.rgb = RGBColor(30, 41, 59)
        year_title_para.paragraph_format.space_after = Pt(4)
        
        # Subtitle
        subtitle_para = content_cell.add_paragraph()
        subtitle_run = subtitle_para.add_run(event["subtitle"])
        subtitle_run.font.size = Pt(9)
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
        subtitle_para.paragraph_format.space_after = Pt(6)
        
        # Items (bullet points with bold keywords)
        for item in event["items"]:
            item_para = content_cell.add_paragraph()
            item_para.paragraph_format.left_indent = Inches(0.25)
            item_para.paragraph_format.space_after = Pt(3)
            
            # Remove ** markers and format with bold keywords
            clean_item = item.replace("**", "")
            
            # Find keyword (text before colon)
            parts = item.split("**")
            bullet_run = item_para.add_run("• ")
            bullet_run.font.size = Pt(10)
            bullet_run.font.color.rgb = RGBColor(55, 65, 81)
            
            # Add text with bold keywords
            for i, part in enumerate(parts):
                if part:
                    text_run = item_para.add_run(part)
                    text_run.font.size = Pt(10)
                    text_run.font.color.rgb = RGBColor(55, 65, 81)
                    # Bold keywords (odd indices after split on **)
                    if i % 2 == 1:
                        text_run.font.bold = True
        
        # Add spacing after event
        doc.add_paragraph().paragraph_format.space_after = Pt(10)
    
    # Footer note
    doc.add_paragraph().paragraph_format.space_before = Pt(20)
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("➡ Commitment extends to Net Zero Operational Emissions by 2030 and Financed Emissions by 2050.")
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(55, 65, 81)
    
    return doc


def create_table():
    """Create roadmap as compact vertical layout for insertion into cell (half-page layout)"""
    doc = Document()
    
    for idx, event in enumerate(ROADMAP_EVENTS):
        # Year and title
        year_title_para = doc.add_paragraph()
        year_title_text = f"{event['year']}: {event['title']}"
        year_title_run = year_title_para.add_run(year_title_text)
        year_title_run.font.size = Pt(11)
        year_title_run.font.bold = True
        year_title_run.font.color.rgb = RGBColor(30, 41, 59)
        year_title_para.paragraph_format.space_before = Pt(8) if idx > 0 else Pt(0)
        year_title_para.paragraph_format.space_after = Pt(2)
        
        # Apply left border
        pPr = year_title_para._element.get_or_add_pPr()
        left_border = parse_xml(r'<w:pBdr {}><w:left w:val="single" w:sz="480" w:space="4" w:color="1F2937"/></w:pBdr>'.format(nsdecls('w')))
        existing_borders = pPr.xpath('.//w:pBdr')
        for b in existing_borders:
            pPr.remove(b)
        pPr.append(left_border)
        year_title_para.paragraph_format.left_indent = Inches(0.2)
        
        # Subtitle
        subtitle_para = doc.add_paragraph()
        subtitle_run = subtitle_para.add_run(event["subtitle"])
        subtitle_run.font.size = Pt(8)
        subtitle_run.font.color.rgb = RGBColor(102, 102, 102)
        subtitle_para.paragraph_format.space_after = Pt(4)
        subtitle_para.paragraph_format.left_indent = Inches(0.2)
        
        # Items
        for item in event["items"]:
            item_para = doc.add_paragraph()
            item_para.paragraph_format.left_indent = Inches(0.35)
            item_para.paragraph_format.space_after = Pt(2)
            
            # Add text with bold keywords
            parts = item.split("**")
            bullet_run = item_para.add_run("• ")
            bullet_run.font.size = Pt(8)
            bullet_run.font.color.rgb = RGBColor(55, 65, 81)
            
            for i, part in enumerate(parts):
                if part:
                    text_run = item_para.add_run(part)
                    text_run.font.size = Pt(8)
                    text_run.font.color.rgb = RGBColor(55, 65, 81)
                    if i % 2 == 1:
                        text_run.font.bold = True
    
    return doc


def create_roadmap_table():
    """Alias for create_table() - for backward compatibility"""
    return create_table()


if __name__ == "__main__":
    import os
    import time
    
    # Test: save as standalone document (full timeline layout)
    doc = create_roadmap()
    timestamp = int(time.time())
    filename1 = f"preview_esg_roadmap_{timestamp}.docx"
    doc.save(filename1)
    
    # Test: save compact version (for half-page layout)
    compact_doc = create_table()
    filename2 = f"preview_esg_roadmap_compact_{timestamp}.docx"
    compact_doc.save(filename2)
    
    print("Timeline flowchart generated successfully!")
    print(f"- {filename1}: Full timeline with vertical line, circular markers, and styled boxes")
    print(f"- {filename2}: Compact vertical layout for half-page")
