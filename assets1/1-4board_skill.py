"""Board skill coverage table for Word reports."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

SKILL_ROWS = [
    {"skill": "Climate / Sustainability", "current": 85, "target": 100, "notes": "TCFD oversight, transition plan review"},
    {"skill": "AI / Data Ethics", "current": 60, "target": 100, "notes": "Responsible AI, data governance, privacy"},
    {"skill": "Cyber Security", "current": 65, "target": 100, "notes": "Resilience, incident response, supplier security"},
    {"skill": "Human Capital Mgmt", "current": 75, "target": 100, "notes": "Succession, DEI, workforce transformation"},
    {"skill": "Financial / Risk", "current": 100, "target": 100, "notes": "Capital allocation, risk appetite, controls"},
]


def _shade(cell, color_hex):
    shading = parse_xml(r'<w:shd {} w:val="clear" w:color="auto" w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading)


def _write(cell, text, *, bold=False, size=10, color=(31, 41, 55), align_center=False):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(2)
        paragraph_format.line_spacing = 1.05
        if align_center:
            paragraph.alignment = 1
        for run in paragraph.runs:
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = RGBColor(*color)


def create_table():
    doc = Document()
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"

    headers = ["Skill Domain", "Current Coverage (%)", "Target Coverage (%)", "Key Governance Notes"]
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        _write(header_cells[idx], header, bold=True, size=11, color=(255, 255, 255), align_center=True)
        _shade(header_cells[idx], "0F4C81")

    for row in SKILL_ROWS:
        cells = table.add_row().cells
        _write(cells[0], row["skill"], bold=True)
        _shade(cells[0], "E0F2FF")

        _write(cells[1], f"{row['current']}%", bold=True, align_center=True, color=(12, 74, 110))
        _shade(cells[1], "F1F5F9")

        _write(cells[2], f"{row['target']}%", bold=True, align_center=True, color=(22, 101, 52))
        _shade(cells[2], "ECFDF5")

        _write(cells[3], row["notes"], bold=False)

    return table


if __name__ == "__main__":
    table = create_table()
    doc = Document()
    doc._element.body.append(table._element)
    doc.save("preview_board_skill.docx")
