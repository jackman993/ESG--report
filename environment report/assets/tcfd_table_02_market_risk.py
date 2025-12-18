from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


def create_market_risk_table(doc, is_first_page=False):
    """市場風險表格 - 統一風格"""
    
    if not is_first_page:
        doc.add_page_break()
    
    # 橫向設定
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # 建立表格 (4行: 主標題 + 欄位標題 + 2資料)
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    
    # ========== Row 0: 分割主標題 ==========
    title_row = table.rows[0]
    
    # 左側綠色
    left_cell = table.cell(0, 0).merge(table.cell(0, 2))
    left_cell.text = 'Climate-Related Risks'
    para = left_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_background(left_cell, '2F5233')
    _set_cell_vertical_alignment(left_cell, 'center')
    
    # 右側灰色
    right_cell = table.cell(0, 3).merge(table.cell(0, 5))
    right_cell.text = 'Financial Impacts'
    para = right_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_background(right_cell, '808080')
    _set_cell_vertical_alignment(right_cell, 'center')
    
    _set_row_height(title_row, Inches(0.5))
    
    # ========== Row 1: 欄位標題 ==========
    headers = ['Type', 'Climate Change\nRelated Risk', 'Impact\nPeriod',
               'Description of Risk Content',
               'Potential Impact on Business,\nStrategies & Finance',
               'Adaption & Responsive\nActions']
    
    header_row = table.rows[1]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_background(cell, '808080')
        _set_cell_vertical_alignment(cell, 'center')
    
    # 欄寬
    widths = [Inches(0.9), Inches(1.2), Inches(0.8), 
              Inches(2.8), Inches(2.8), Inches(2.1)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    
    # ========== Row 2: Market Risk ==========
    row2 = table.rows[2]
    
    # Type
    cell = row2.cells[0]
    cell.text = 'Transformation\nRisk'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_background(cell, '8B9D83')
    _set_cell_vertical_alignment(cell, 'center')
    
    # Climate Risk
    cell = row2.cells[1]
    cell.text = 'Market'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')
    
    # Period
    cell = row2.cells[2]
    cell.text = 'Long-term'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')
    
    # Description
    market_desc = [
        'Consumer environmental awareness drives purchasing decisions, with 65-75% of buyers considering product carbon footprint in procurement.',
        'Market demand for low-carbon products accelerates, requiring rapid product portfolio transition and lifecycle assessment capabilities.',
        'Competitors gaining market share through verified environmental credentials create pressure to establish comparable sustainability positioning.'
    ]
    _add_bullets(row2.cells[3], market_desc, '2F5233')
    
    # Impact
    market_impact = [
        'Revenue at risk: $24-34M annually from products unable to meet evolving environmental standards (approximately 18-22% of current portfolio).',
        'Market share erosion of 5-8 percentage points to competitors with stronger sustainability credentials in key segments.',
        'Product development costs increase $300-520K to integrate carbon footprint assessment and eco-design principles across portfolio.'
    ]
    _add_bullets(row2.cells[4], market_impact, '2F5233')
    
    # Actions
    market_actions = [
        'Launched Carbon Management Platform in 2021, providing customers with product-level carbon footprint reports verified through quarterly ESG Committee reviews.',
        'Distribute low-carbon products strategically, assisting channel partners with sustainability marketing; implementing transition requests to carbon-reduction technologies.',
        'Established transparent ESG disclosure through annual sustainability reports and real-time stakeholder communication channels.'
    ]
    _add_bullets(row2.cells[5], market_actions, '2F5233')
    
    # ========== Row 3: Consumer Awareness ==========
    row3 = table.rows[3]
    
    # Type (淺綠背景)
    _set_cell_background(row3.cells[0], '8B9D83')
    
    # Climate Risk
    cell = row3.cells[1]
    cell.text = "Consumer's\nsustainability\nawareness"
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')
    
    # Period
    cell = row3.cells[2]
    cell.text = 'Long-term'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')
    
    # Description
    consumer_desc = [
        'Rising labor and operational costs associated with conducting lifecycle assessments (LCA) impact revenues and profitability margins.',
        'Cloud services and low-carbon product demand increases, requiring accelerated market transition and technology investment.',
        'Supply chain partners increasingly require carbon disclosure and reduction commitments as precondition for contract renewals.'
    ]
    _add_bullets(row3.cells[3], consumer_desc, '2F5233')
    
    # Impact
    consumer_impact = [
        'LCA study costs accumulated to approximately $340,000 since 2010, with ongoing annual expenses of $85,000-120,000 for portfolio coverage.',
        'Agency and product development costs for low-carbon alternatives estimated at $280-420K to maintain competitive positioning.',
        'Market competitiveness requires continuous investment in carbon reduction practices, with sustainability becoming a primary purchase criterion affecting 40-55% of procurement decisions.'
    ]
    _add_bullets(row3.cells[4], consumer_impact, '2F5233')
    
    # Actions
    consumer_actions = [
        'Implemented carbon reduction practices with transparent disclosure through ESG reports and proactive stakeholder communications.',
        'Developed low-carbon product lines addressing green market competition, supporting distributors with sustainability-focused sales materials.',
        'Established carbon inventory systems enabling data-driven reduction strategies and facilitating customer environmental compliance requirements.'
    ]
    _add_bullets(row3.cells[5], consumer_actions, '2F5233')
    
    # 合併 Type 欄位
    table.cell(2, 0).merge(table.cell(3, 0))
    
    return doc


def _set_cell_background(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)


def _set_cell_vertical_alignment(cell, align='center'):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def _set_row_height(row, height):
    tr = row._element
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height.pt * 20)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def _add_bullets(cell, items, bullet_color):
    cell.text = ''
    for item in items:
        para = cell.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.12)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.line_spacing = 1.0
        
        run = para.add_run('▲ ')
        run.font.size = Pt(8)
        color_rgb = RGBColor(
            int(bullet_color[:2], 16),
            int(bullet_color[2:4], 16),
            int(bullet_color[4:], 16)
        )
        run.font.color.rgb = color_rgb
        
        run = para.add_run(item)
        run.font.size = Pt(8)
        run.font.name = 'Arial'


# ============ 執行 ============
if __name__ == "__main__":
    doc = Document()
    create_market_risk_table(doc, is_first_page=True)
    doc.save('market_risk_table.docx')
    
    print("✓ Market Risk 表格生成完成!")
    print("✓ 包含: Market + Consumer Awareness")
    print("✓ 統一綠灰分割風格")
    
    try:
        from google.colab import files
        files.download('market_risk_table.docx')
    except:
        print("(非 Colab 環境)")