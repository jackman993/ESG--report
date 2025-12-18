from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


def create_physical_risk_table(doc, is_first_page=False):
    """實體風險表格 - 綠灰分割版"""
    
    if not is_first_page:
        doc.add_page_break()
    
    # 橫向設定
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # 建立表格 (4行: 主標題 + 欄位標題 + 2資料)
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    
    # ========== Row 0: 分割主標題 ==========
    title_row = table.rows[0]
    
    # 左側綠色
    left_cell = table.cell(0, 0).merge(table.cell(0, 2))
    left_cell.text = 'Climate-Related Risks'
    para = left_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_background(left_cell, '2F5233')
    _set_cell_vertical_alignment(left_cell, 'center')
    
    # 右側灰色
    right_cell = table.cell(0, 3).merge(table.cell(0, 5))
    right_cell.text = 'Financial Impacts'
    para = right_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_background(right_cell, '808080')
    _set_cell_vertical_alignment(right_cell, 'center')
    
    _set_row_height(title_row, Inches(0.5))
    
    # ========== Row 1: 欄位標題 ==========
    headers = ['Type', 'Climate Change\nRelated Risk', 'Impact\nPeriod',
               'Description of Risk Content',
               'Potential Impact on Business,\nStrategies & Finance',
               'Adaption & Responsive\nActions']
    
    header_row = table.rows[1]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_background(cell, '808080')
        _set_cell_vertical_alignment(cell, 'center')
    
    # 欄寬
    widths = [Inches(0.9), Inches(1.2), Inches(0.8), 
              Inches(2.8), Inches(2.8), Inches(2.1)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    
    # ========== Row 2: Extreme Weather Events ==========
    row2 = table.rows[2]
    
    # Type
    cell = row2.cells[0]
    cell.text = 'Physical\nRisk'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_background(cell, '8B9D83')
    _set_cell_vertical_alignment(cell, 'center')
    
    # Climate Risk
    cell = row2.cells[1]
    cell.text = 'Extreme\nWeather\nEvents'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')
    
    # Period
    cell = row2.cells[2]
    cell.text = 'Short-term\nto\nmedium-term'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')
    
    # Description
    weather_desc = [
        'Increasing frequency and intensity of extreme weather events (typhoons, floods, heatwaves) threaten operational continuity, with climate models projecting 30-45% increase in severe weather incidents by 2030.',
        'Infrastructure vulnerabilities include flooding risks, power outages lasting 12-48 hours, and supply chain disruptions affecting material delivery and product shipment.',
        'Geographic concentration in climate-vulnerable regions exposes 55-70% of production capacity to physical climate risks requiring comprehensive resilience measures.'
    ]
    _add_bullets(row2.cells[3], weather_desc, '2F5233')
    
    # Impact
    weather_impact = [
        'Production downtime costs estimated at $18,000-35,000 per day for facility disruptions, with weather events historically causing 3-7 days of operational interruption annually.',
        'Emergency response expenses including facility repairs, equipment replacement, and expedited logistics average $120,000-280,000 per significant weather event.',
        'Insurance premiums increasing 8-15% annually for climate coverage, with deductibles rising and certain extreme weather protections becoming cost-prohibitive.'
    ]
    _add_bullets(row2.cells[4], weather_impact, '2F5233')
    
    # Actions
    weather_actions = [
        'Deployed weather monitoring systems and early warning applications enabling 24-48 hour advance preparation, integrated with business continuity protocols for pre-positioning resources and adjusting production schedules.',
        'Invested $450,000-720,000 in resilience infrastructure including enhanced drainage systems, flood barriers, backup diesel generators (48-72 hours capacity), 5kW solar power installations with 200kWh battery storage systems ensuring continuous operation during grid outages and extreme weather events.',
        'Established water management infrastructure with storm water retention capacity, improved drainage systems, and flood prevention protocols, while building emergency reserves equivalent to 1-2 weeks operational needs.'
    ]
    _add_bullets(row2.cells[5], weather_actions, '2F5233')
    
    # ========== Row 3: Resource Availability ==========
    row3 = table.rows[3]
    
    # Type (淺綠背景)
    _set_cell_background(row3.cells[0], '8B9D83')
    
    # Climate Risk
    cell = row3.cells[1]
    cell.text = 'Resource\nAvailability'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')
    
    # Period
    cell = row3.cells[2]
    cell.text = 'Medium-term\nto\nlong-term'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')
    
    # Description
    resource_desc = [
        'Water stress and energy supply constraints in operating regions threaten production continuity, with regional water scarcity projected to intensify 25-40% over next decade.',
        'Competition for limited resources drives regulatory restrictions potentially prioritizing residential and agricultural needs over industrial users during shortage periods.',
            ]
    _add_bullets(row3.cells[3], resource_desc, '2F5233')
    
    # Impact
    resource_impact = [
        'Water and energy cost escalation estimated at 12-22% over 5-year horizon due to resource scarcity, adding $250,000-480,000 to annual operational expenses.',
        'Production capacity at risk: potential 15-25% output reduction if unable to secure adequate water allocations or experiencing extended power interruptions.',
            ]
    _add_bullets(row3.cells[4], resource_impact, '2F5233')
    
    # Actions
    resource_actions = [
        
      'Deployed real-time resource monitoring systems tracking water availability, energy grid status, and weather forecasts through integrated applications, enabling proactive production planning and resource procurement adjustments.'
    ]
    _add_bullets(row3.cells[5], resource_actions, '2F5233')
    
    # 合併 Type 欄位
    table.cell(2, 0).merge(table.cell(3, 0))
    
    return doc


def _set_cell_background(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)


def _set_cell_vertical_alignment(cell, align='center'):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def _set_row_height(row, height):
    tr = row._element
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height.pt * 20)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def _add_bullets(cell, items, bullet_color):
    cell.text = ''
    for item in items:
        para = cell.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.12)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.line_spacing = 1.0
        
        run = para.add_run('▲ ')
        run.font.size = Pt(8)
        color_rgb = RGBColor(
            int(bullet_color[:2], 16),
            int(bullet_color[2:4], 16),
            int(bullet_color[4:], 16)
        )
        run.font.color.rgb = color_rgb
        
        run = para.add_run(item)
        run.font.size = Pt(8)
        run.font.name = 'Arial'


# ============ 執行 ============
if __name__ == "__main__":
    doc = Document()
    create_physical_risk_table(doc, is_first_page=True)
    doc.save('tcfd_table_03_physical_risk.docx')
    
    print("✓ Physical Risk 表格生成完成!")
    print("✓ 包含: Extreme Weather Events + Resource Availability")
    print("✓ 實務措施: 天氣APP、5kW太陽能、200kWh儲能、水利系統")
    
    try:
        from google.colab import files
        files.download('tcfd_table_03_physical_risk.docx')
    except:
        print("(非 Colab 環境)")