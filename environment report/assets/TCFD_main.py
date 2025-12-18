from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_climate_risk_timeline_english(doc, is_first_page=False):
    """Climate Risk Timeline Table - English Version"""

    if not is_first_page:
        doc.add_page_break()

    # 橫向設定
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Cm(1.0) + Inches(0.4)   # 原 0.4" + 1cm 留白
    section.right_margin = Cm(1.0) + Inches(0.4)  # 原 0.4" + 1cm 留白
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # 添加標題
    title = doc.add_heading('Climate-Related Risk Timeline', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(16)  # 18 * 0.85 ≈ 15.3 → 16
    title_run.font.color.rgb = RGBColor(47, 82, 51)  # 深綠色

    doc.add_paragraph()  # 空行

    # 建立表格 (15行4列: 表頭 + 14個風險項目)
    table = doc.add_table(rows=15, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # ========== Row 0: 表頭 ==========
    header_row = table.rows[0]

    # 合併表頭
    header_cell1 = table.cell(0, 0).merge(table.cell(0, 1))
    header_cell1.text = 'Risk Events'
    para = header_cell1.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(12)  # 14 * 0.85 ≈ 11.9 → 12
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_background(header_cell1, '4a7c59')  # 深綠色
    _set_cell_vertical_alignment(header_cell1, 'center')

    header_cell2 = table.cell(0, 2).merge(table.cell(0, 3))
    header_cell2.text = 'Timeline'
    para = header_cell2.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(12)  # 14 * 0.85 ≈ 11.9 → 12
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_background(header_cell2, '6b9475')  # 中綠色
    _set_cell_vertical_alignment(header_cell2, 'center')

    _set_row_height(header_row, Inches(0.51))  # 0.6 * 0.85 = 0.51

    # 風險數據 - 英文版
    risks_data = [
        ('Transition\nRisk', '1', 'GHG Emission Cost Increase - Carbon Pricing Mechanism', '2025', 'Short-term'),
        ('', '2-1', 'GHG Emission Cost Increase - Carbon Border Adjustment (CCA)', '2025', 'Short-term'),
        ('', '2-2', 'GHG Emission Cost Increase - Carbon Border Adjustment (CBAM)', '2026', 'Medium-term'),
        ('', '3-1', 'Sustainability Requirements & Regulations Increase - Other Industries', '2030', 'Medium-term'),
        ('', '4', 'Investment Risk Increase for High-Carbon Assets', '2030', 'Medium-term'),
        ('', '5', 'Transition Cost to Low-Carbon Economy', '2028', 'Medium-term'),
        ('', '6-1', 'Raw Material Cost Increase (Infrastructure Materials)', '2027', 'Medium-term'),
        ('', '3-2', 'Sustainability Requirements & Regulations Increase - Automotive', '2035', 'Long-term'),
        ('', '6-2', 'Raw Material Cost Increase (Food)', '2050', 'Long-term'),
        ('Physical\nRisk', '7', 'Extreme Weather Events Increase - Typhoons', '2050', 'Long-term'),
        ('', '8', 'Extreme Weather Events Increase - Heavy Rain', '2050', 'Long-term'),
        ('', '9', 'Extreme Weather Events Increase - Drought', '2050', 'Long-term'),
        ('', '10', 'Average Temperature Rise', '2050', 'Long-term'),
        ('', '11', 'Sea Level Rise', '2050', 'Long-term')
    ]

    # 設定欄寬 (縮小至 85%)
    table.columns[0].width = Inches(1.02)  # 1.2 * 0.85 = 1.02  類別
    table.columns[1].width = Inches(0.68)  # 0.8 * 0.85 = 0.68  編號
    table.columns[2].width = Inches(5.1)   # 6.0 * 0.85 = 5.1   風險事件
    table.columns[3].width = Inches(2.13)  # 2.5 * 0.85 = 2.125 時間軸

    # 填入數據
    for i, (category, number, risk_event, year, period) in enumerate(risks_data):
        row = table.rows[i + 1]

        # 類別欄
        if category:  # 只在第一個項目顯示類別
            cell = row.cells[0]
            cell.text = category
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.font.bold = True
            run.font.size = Pt(9)  # 11 * 0.85 ≈ 9.35 → 9
            run.font.name = 'Arial'
            run.font.color.rgb = RGBColor(255, 255, 255)
            if 'Transition' in category:
                _set_cell_background(cell, '8B9D83')  # 淺綠色
            else:
                _set_cell_background(cell, '9BA89E')  # 淺灰綠
            _set_cell_vertical_alignment(cell, 'center')

        # 編號欄
        cell = row.cells[1]
        cell.text = number
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.size = Pt(9)  # 10 * 0.85 = 8.5 → 9
        run.font.name = 'Arial'
        _set_cell_background(cell, 'F5F5F5')  # 淺灰色
        _set_cell_vertical_alignment(cell, 'center')

        # 風險事件欄
        cell = row.cells[2]
        cell.text = risk_event
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = para.runs[0]
        run.font.size = Pt(9)  # 10 * 0.85 = 8.5 → 9
        run.font.name = 'Arial'
        _set_cell_vertical_alignment(cell, 'center')

        # 時間軸欄
        cell = row.cells[3]
        cell.text = f'{year}    {period}'
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.size = Pt(9)  # 10 * 0.85 = 8.5 → 9
        run.font.name = 'Arial'

        # 根據時期設定背景色
        if 'Short-term' in period:
            _set_cell_background(cell, 'E8F5E8')  # 很淺的綠色
        elif 'Medium-term' in period:
            _set_cell_background(cell, 'D4E6D4')  # 淺綠色
        else:  # Long-term
            _set_cell_background(cell, 'C0D6C0')  # 中淺綠色

        _set_cell_vertical_alignment(cell, 'center')
        _set_row_height(row, Inches(0.34))  # 0.4 * 0.85 = 0.34

    # 合併轉型風險類別欄位 (行1-9，共9項)
    table.cell(1, 0).merge(table.cell(9, 0))
    # 合併實體風險類別欄位 (行10-14，共5項)
    table.cell(10, 0).merge(table.cell(14, 0))

    return doc


def _set_cell_background(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)


def _set_cell_vertical_alignment(cell, align='center'):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def _set_row_height(row, height):
    tr = row._element
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    # 將 EMU (Inches 返回值) 轉換為 twips (Word 行高單位)
    # 1 inch = 914400 EMU = 1440 twips, 所以 twips = EMU / 635
    trHeight.set(qn('w:val'), str(int(height / 635)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


# ============ 執行 ============
if __name__ == "__main__":
    doc = Document()
    create_climate_risk_timeline_english(doc, is_first_page=True)
    doc.save('climate_risk_timeline_english.docx')

    print("✓ Climate Risk Timeline Table (English Version) Generated!")
    print("✓ Design Features:")
    print("  - Green gradient color scheme")
    print("  - Transition Risk vs Physical Risk classification")
    print("  - Short-term/Medium-term/Long-term timeline")
    print("  - Professional layout design")

    try:
        from google.colab import files
        files.download('climate_risk_timeline_english.docx')
    except:
        print("(Non-Colab environment)")