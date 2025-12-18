from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


def create_resource_efficiency_table(doc, is_first_page=False):
    """資源效率機會表格 - Resource Efficiency"""

    if not is_first_page:
        doc.add_page_break()

    # 橫向設定
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # 建立表格 (3行: 主標題 + 欄位標題 + 1資料)
    table = doc.add_table(rows=3, cols=6)
    table.style = 'Table Grid'

    # ========== Row 0: 分割主標題 ==========
    title_row = table.rows[0]

    # 左側綠色
    left_cell = table.cell(0, 0).merge(table.cell(0, 2))
    left_cell.text = 'Climate-Related Opportunities'
    para = left_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_background(left_cell, '2F5233')
    _set_cell_vertical_alignment(left_cell, 'center')

    # 右側灰色
    right_cell = table.cell(0, 3).merge(table.cell(0, 5))
    right_cell.text = 'Financial Impacts'
    para = right_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_background(right_cell, '808080')
    _set_cell_vertical_alignment(right_cell, 'center')

    _set_row_height(title_row, Inches(0.5))

    # ========== Row 1: 欄位標題 ==========
    headers = ['Type', 'Climate Change\nRelated Opportunity', 'Impact\nPeriod',
               'Description of Opportunity Content',
               'Potential Impact on Business,\nStrategies & Finance',
               'Implementation &\nResponsive Actions']

    header_row = table.rows[1]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_background(cell, '808080')
        _set_cell_vertical_alignment(cell, 'center')

    # 欄寬
    widths = [Inches(0.9), Inches(1.2), Inches(0.8), 
              Inches(2.8), Inches(2.8), Inches(2.1)]
    for i, width in enumerate(widths):
        table.columns[i].width = width

    # ========== Row 2: Resource Efficiency ==========
    row2 = table.rows[2]

    # Type
    cell = row2.cells[0]
    cell.text = 'Resource\nEfficiency'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_background(cell, '8B9D83')
    _set_cell_vertical_alignment(cell, 'center')

    # Climate Opportunity
    cell = row2.cells[1]
    cell.text = 'Energy-saving\nbenefit'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')

    # Period
    cell = row2.cells[2]
    cell.text = 'Short-term'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')

    # Description
    resource_desc = [
        'Intelligent energy management systems deployment across all facilities utilizing AI-driven predictive analytics to optimize power consumption patterns and eliminate operational inefficiencies.',
        'Digital transformation of energy monitoring and optimization processes through IoT sensor networks, real-time data analytics, and automated control systems for comprehensive energy visibility.',
        'Circular economy integration with waste-to-energy conversion initiatives, heat recovery systems, and material reuse programs creating sustainable value chains.'
    ]
    _add_bullets(row2.cells[3], resource_desc, '2F5233')

    # Impact
    resource_impact = [
        'AI-driven energy optimization reduces operational costs by $300,000-900,000 annually through predictive load management, demand response optimization, and elimination of energy waste across facilities.',
        'Energy digitalization generates $250,000-750,000 savings per year via real-time monitoring, automated controls, and data-driven efficiency improvements reducing overall energy consumption by 25-35%.',
        'Circular value networks create $200,000-600,000 new revenue from waste heat recovery, material reuse programs, and strategic partnerships converting waste streams into profitable secondary resources.'
    ]
    _add_bullets(row2.cells[4], resource_impact, '2F5233')

    # Actions
    resource_actions = [
        'Deploy comprehensive IoT sensors and AI analytics infrastructure for predictive energy optimization achieving 25-40% efficiency improvements through machine learning algorithms, automated load balancing, and real-time performance optimization.',
        'Establish integrated energy data platform connecting all operational systems for real-time performance tracking, predictive maintenance scheduling, and automated energy management with centralized dashboard monitoring.',
        'Build strategic partnerships for circular economy ecosystems converting waste streams into profitable resources through supplier collaboration, waste-to-energy projects, and innovative material recovery programs generating additional revenue channels.'
    ]
    _add_bullets(row2.cells[5], resource_actions, '2F5233')

    return doc


def _set_cell_background(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)


def _set_cell_vertical_alignment(cell, align='center'):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def _set_row_height(row, height):
    tr = row._element
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height.pt * 20)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def _add_bullets(cell, items, bullet_color):
    cell.text = ''
    for item in items:
        para = cell.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.12)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.line_spacing = 1.0

        run = para.add_run('▲ ')
        run.font.size = Pt(8)
        color_rgb = RGBColor(
            int(bullet_color[:2], 16),
            int(bullet_color[2:4], 16),
            int(bullet_color[4:], 16)
        )
        run.font.color.rgb = color_rgb

        run = para.add_run(item)
        run.font.size = Pt(8)
        run.font.name = 'Arial'


# ============ 執行 ============
if __name__ == "__main__":
    doc = Document()
    create_resource_efficiency_table(doc, is_first_page=True)
    doc.save('tcfd_table_05_resource_efficiency.docx')

    print("✓ Resource Efficiency 表格生成完成!")
    print("✓ 包含能源效率機會:")
    print("  - AI驅動能源優化")
    print("  - 數位化監控系統")
    print("  - 循環經濟整合")
    print("  - 金額範圍: $200K-900K")

    try:
        from google.colab import files
        files.download('tcfd_table_05_resource_efficiency.docx')
    except:
        print("(非 Colab 環境)")