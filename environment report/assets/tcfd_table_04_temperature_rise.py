from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


def create_temperature_risk_table(doc, is_first_page=False):
    """溫度上升風險表格 - 含建築設計"""
    
    if not is_first_page:
        doc.add_page_break()
    
    # 橫向設定
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # 建立表格 (3行: 主標題 + 欄位標題 + 1資料)
    table = doc.add_table(rows=3, cols=6)
    table.style = 'Table Grid'
    
    # ========== Row 0: 分割主標題 ==========
    title_row = table.rows[0]
    
    # 左側綠色
    left_cell = table.cell(0, 0).merge(table.cell(0, 2))
    left_cell.text = 'Climate-Related Risks'
    para = left_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_background(left_cell, '2F5233')
    _set_cell_vertical_alignment(left_cell, 'center')
    
    # 右側灰色
    right_cell = table.cell(0, 3).merge(table.cell(0, 5))
    right_cell.text = 'Financial Impacts'
    para = right_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_background(right_cell, '808080')
    _set_cell_vertical_alignment(right_cell, 'center')
    
    _set_row_height(title_row, Inches(0.5))
    
    # ========== Row 1: 欄位標題 ==========
    headers = ['Type', 'Climate Change\nRelated Risk', 'Impact\nPeriod',
               'Description of Risk Content',
               'Potential Impact on Business,\nStrategies & Finance',
               'Adaption & Responsive\nActions']
    
    header_row = table.rows[1]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_background(cell, '808080')
        _set_cell_vertical_alignment(cell, 'center')
    
    # 欄寬
    widths = [Inches(0.9), Inches(1.2), Inches(0.8), 
              Inches(2.8), Inches(2.8), Inches(2.1)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    
    # ========== Row 2: Temperature Rise ==========
    row2 = table.rows[2]
    
    # Type
    cell = row2.cells[0]
    cell.text = 'Physical\nRisk'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_background(cell, '8B9D83')
    _set_cell_vertical_alignment(cell, 'center')
    
    # Climate Risk
    cell = row2.cells[1]
    cell.text = 'Increase of\nannual\ntemperature'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')
    
    # Period
    cell = row2.cells[2]
    cell.text = 'Medium-term\nto\nlong-term'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')
    
    # Description
    temp_desc = [
        'Rising ambient temperatures reduce equipment operational efficiency and accelerate asset degradation, with manufacturing equipment designed for 15-25°C optimal performance facing stress at sustained temperatures above 30°C.',
        'Heat stress impacts workforce productivity and safety, particularly for non-climate-controlled areas, with studies showing 10-15% productivity decline when temperatures exceed 28°C and increased occupational health risks.',
        'Cooling energy demand escalating 25-35% over next decade to maintain operational temperature requirements, with peak summer loads creating grid reliability concerns and cost pressures.'
    ]
    _add_bullets(row2.cells[3], temp_desc, '2F5233')
    
    # Impact
    temp_impact = [
        'Equipment maintenance costs increasing 8-12% annually due to thermal stress-induced failures, with critical machinery requiring enhanced cooling systems estimated at $300,000-550,000 investment.',
        'Labor productivity losses estimated at $80,000-150,000 annually from heat-related work slowdowns, absenteeism, and required safety protocols including extended breaks and shift adjustments.',
        'Energy expenses for climate control rising $180,000-320,000 per year, with HVAC systems representing 35-45% of facility energy consumption and requiring capacity upgrades to meet temperature requirements.'
    ]
    _add_bullets(row2.cells[4], temp_impact, '2F5233')
    
    # Actions
    temp_actions = [
        'Invested $420,000-680,000 in thermal management infrastructure including high-efficiency HVAC systems, low-emissivity (Low-E) glass windows for solar heat reduction, open-plan office design promoting natural cross-ventilation, vertical greenery systems (living walls) providing passive cooling and air quality improvement, and building envelope upgrades to maintain optimal operating temperatures (18-26°C, 40-60% humidity).',
        'Implemented adaptive work policies including flexible scheduling during extreme heat periods, heat stress prevention training, strategically placed hydration stations, temperature monitoring systems with automated alerts, and enhanced break areas with climate-controlled rest zones to maintain productivity and worker safety.',
        'Deployed intelligent energy management systems optimizing cooling loads through predictive algorithms, thermal energy storage during off-peak hours, integration with weather forecasting for proactive pre-cooling before heat waves, and automated shading systems responsive to solar radiation, achieving 15-22% reduction in cooling energy intensity.'
    ]
    _add_bullets(row2.cells[5], temp_actions, '2F5233')
    
    return doc


def _set_cell_background(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)


def _set_cell_vertical_alignment(cell, align='center'):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def _set_row_height(row, height):
    tr = row._element
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height.pt * 20)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def _add_bullets(cell, items, bullet_color):
    cell.text = ''
    for item in items:
        para = cell.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.12)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.line_spacing = 1.0
        
        run = para.add_run('▲ ')
        run.font.size = Pt(8)
        color_rgb = RGBColor(
            int(bullet_color[:2], 16),
            int(bullet_color[2:4], 16),
            int(bullet_color[4:], 16)
        )
        run.font.color.rgb = color_rgb
        
        run = para.add_run(item)
        run.font.size = Pt(8)
        run.font.name = 'Arial'


# ============ 執行 ============
if __name__ == "__main__":
    doc = Document()
    create_temperature_risk_table(doc, is_first_page=True)
    doc.save('tcfd_table_04_temperature_rise.docx')
    
    print("✓ Temperature Rise 表格生成完成!")
    print("✓ 包含建築設計元素:")
    print("  - Low-E 玻璃隔熱")
    print("  - 開放式對流設計")
    print("  - 垂直綠化系統")
    print("  - 智能能源管理")
    
    try:
        from google.colab import files
        files.download('tcfd_table_04_temperature_rise.docx')
    except:
        print("(非 Colab 環境)")