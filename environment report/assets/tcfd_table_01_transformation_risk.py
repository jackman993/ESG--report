from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


def create_split_header_table(doc, is_first_page=True):
    """綠灰分割表頭版本"""
    
    if not is_first_page:
        doc.add_page_break()
    
    # 橫向設定
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    
    # 建立表格 (4行: 1主標題 + 1欄位標題 + 2資料)
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    
    # ========== Row 0: 分割主標題 ==========
    title_row = table.rows[0]
    
    # 左側綠色 (合併前3欄)
    left_cell = table.cell(0, 0).merge(table.cell(0, 2))
    left_cell.text = 'Climate-Related Risks'
    para = left_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_background(left_cell, '2F5233')  # 深綠
    _set_cell_vertical_alignment(left_cell, 'center')
    
    # 右側灰色 (合併後3欄)
    right_cell = table.cell(0, 3).merge(table.cell(0, 5))
    right_cell.text = 'Financial Impacts'
    para = right_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    run.font.color.rgb = RGBColor(255, 255, 255)
    _set_cell_background(right_cell, '808080')  # 灰色
    _set_cell_vertical_alignment(right_cell, 'center')
    
    # 設定標題列高度
    _set_row_height(title_row, Inches(0.5))
    
    # ========== Row 1: 欄位標題 (全灰色) ==========
    headers = ['Type', 'Climate Change\nRelated Risk', 'Impact\nPeriod',
               'Description of Risk Content',
               'Potential Impact on Business,\nStrategies & Finance',
               'Adaption & Responsive\nActions']
    
    header_row = table.rows[1]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_background(cell, '808080')  # 灰色
        _set_cell_vertical_alignment(cell, 'center')
    
    # 欄寬
    widths = [Inches(0.9), Inches(1.2), Inches(0.8), 
              Inches(2.8), Inches(2.8), Inches(2.1)]
    for i, width in enumerate(widths):
        table.columns[i].width = width
    
    # ========== Row 2: Policy and Regulation ==========
    row2 = table.rows[2]
    
    # Type
    cell = row2.cells[0]
    cell.text = 'Transformation\nRisk'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_background(cell, '8B9D83')  # 淺綠
    _set_cell_vertical_alignment(cell, 'center')
    
    # Climate Risk
    cell = row2.cells[1]
    cell.text = 'Policy and\nRegulation'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')
    
    # Period
    cell = row2.cells[2]
    cell.text = 'Short-term\nand\nmedium-term'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')
    
    # Description
    policy_desc = [
        'Regulatory pressure accelerates through mandatory emissions reporting and carbon pricing, with penalties of $50K-$2M for non-compliance.',
        'Compliance timelines compressed from 10-year to 3-5-year implementation periods.',
        'Customer procurement now mandates supplier regulatory compliance for 35-40% of revenue base.'
    ]
    _add_bullets(row2.cells[3], policy_desc, '2F5233')
    
    # Impact
    policy_impact = [
        'Direct compliance costs: $1.2-3.5M annually for reporting systems and verification.',
        'Capital expenditure: $8-15M over three years for energy efficiency retrofits.',
        'Contract termination risk if unable to demonstrate compliance; early adoption creates competitive advantage.'
    ]
    _add_bullets(row2.cells[4], policy_impact, '2F5233')
    
    # Actions
    policy_actions = [
        'Quarterly horizon scanning identifies regulations 18-24 months ahead; $2.8M invested in carbon accounting platform.',
        'Cross-functional teams translate requirements into metrics; ESG weighted 15-20% in management incentives.',
        'Active participation in industry coalitions to shape implementation pathways.'
    ]
    _add_bullets(row2.cells[5], policy_actions, '2F5233')
    
    # ========== Row 3: Technology ==========
    row3 = table.rows[3]
    
    # Type (淺綠背景,不合併)
    _set_cell_background(row3.cells[0], '8B9D83')
    
    # Climate Risk
    cell = row3.cells[1]
    cell.text = 'Green product\nand technology'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')
    
    # Period
    cell = row3.cells[2]
    cell.text = 'Long-term'
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.runs[0]
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    _set_cell_vertical_alignment(cell, 'center')
    
    # Description
    tech_desc = [
        'Product development shifts to carbon-performance optimization; 45-60% of sales require ISO 14067 carbon footprint verification.',
        'Manufacturing disruption requires $120-200K capital over 5-7 years for electrification and AI-optimized energy management.',
        'Market bifurcation: early movers capture 8-15% sustainability premiums.'
    ]
    _add_bullets(row3.cells[3], tech_desc, '2F5233')
    
    # Impact
    tech_impact = [
        'R&D reallocation increased from 8% to 22% ($450-680K annually); payback extended to 4-6 years.',
        'Low-carbon materials increase costs 7-18%, compressing margins 2-4 percentage points.',
        'Green segments growing 18-25% CAGR vs 3-5% conventional; potential $8-12M incremental revenue by 2028.'
    ]
    _add_bullets(row3.cells[4], tech_impact, '2F5233')
    
    # Actions
    tech_actions = [
        'Clean Technology Innovation Lab ($5.2M funding) accelerates development cycles from 24 to 14-16 months.',
        'Carbon footprint modeling integrated at each design gate, achieving 12-28% intensity reductions.',
        'Top 30 suppliers (75% spend) required to set science-based targets; joint ventures share R&D risk and IP.'
    ]
    _add_bullets(row3.cells[5], tech_actions, '2F5233')
    
    # 合併 Type 欄位
    table.cell(2, 0).merge(table.cell(3, 0))
    
    return doc


def _set_cell_background(cell, color_hex):
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    cell._element.get_or_add_tcPr().append(shading_elm)


def _set_cell_vertical_alignment(cell, align='center'):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)


def _set_row_height(row, height):
    tr = row._element
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement('w:trHeight')
    trHeight.set(qn('w:val'), str(int(height.pt * 20)))
    trHeight.set(qn('w:hRule'), 'atLeast')
    trPr.append(trHeight)


def _add_bullets(cell, items, bullet_color):
    cell.text = ''
    for item in items:
        para = cell.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.12)
        para.paragraph_format.space_after = Pt(3)
        para.paragraph_format.line_spacing = 1.0
        
        run = para.add_run('▲ ')
        run.font.size = Pt(8)
        color_rgb = RGBColor(
            int(bullet_color[:2], 16),
            int(bullet_color[2:4], 16),
            int(bullet_color[4:], 16)
        )
        run.font.color.rgb = color_rgb
        
        run = para.add_run(item)
        run.font.size = Pt(8)
        run.font.name = 'Arial'


# ============ 執行 ============
if __name__ == "__main__":
    doc = Document()
    create_split_header_table(doc, is_first_page=True)
    doc.save('split_header_table.docx')
    
    print("✓ 綠灰分割表頭生成完成!")
    print("✓ 左側綠色: Climate-Related Risks")
    print("✓ 右側灰色: Financial Impacts")
    
    try:
        from google.colab import files
        files.download('split_header_table.docx')
    except:
        print("(非 Colab 環境)")